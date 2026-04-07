"""
TRICH XUAT HOP DONG CONG CHUNG -> JSON CHO WEB FORM
===================================================
Cach dung:
    python extract_contract.py "hop_dong.docx"
    python extract_contract.py "hop_dong.doc"

Ket qua:
    - In JSON ra man hinh
    - Luu file .json cung thu muc

Cai dat:
    pip install python-docx pywin32
"""

from __future__ import annotations

import json
import os
import re
import sys
import unicodedata
from pathlib import Path
from typing import Iterable


# ============================================================
# CAU HINH MAC DINH — sua tai day cho phu hop van phong
# ============================================================
DEFAULT_CCV = "Phạm Minh Chi"
DEFAULT_THU_KY = "Nguyễn Nhật Minh"
CONTRACT_YEAR = "2026"
CONTRACT_NO_REGEX = re.compile(rf"\b(\d+/{re.escape(CONTRACT_YEAR)}(?:/CCGD)?)\b", re.IGNORECASE)
CONTRACT_KEYWORDS = ("hợp đồng", "hop dong", "hđ", "hd")

NHOM_HD_MAP = {
    "chuyển nhượng": "Chuyển nhượng - Mua bán",
    "mua bán": "Chuyển nhượng - Mua bán",
    "tặng cho": "Tặng, cho tài sản",
    "tặng, cho": "Tặng, cho tài sản",
    "thế chấp": "Cầm cố - Thế chấp - Vay",
    "cầm cố": "Cầm cố - Thế chấp - Vay",
    "vay": "Cầm cố - Thế chấp - Vay",
    "ủy quyền": "Ủy quyền",
    "ủy quền": "Ủy quyền",
    "di chúc": "Di chúc",
    "đặt cọc": "Đặt cọc",
    "thỏa thuận": "Thỏa thuận - Cam kết",
    "cam kết": "Thỏa thuận - Cam kết",
    "góp vốn": "Góp vốn - Hợp tác",
    "hợp tác": "Góp vốn - Hợp tác",
    "thừa kế": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "khai nhận": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "phân chia": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "từ chối nhận di sản": "Từ chối nhận di sản thừa kế",
    "phụ lục": "Phụ lục hợp đồng - văn bản sửa đổi",
    "thuê": "Thuê - Mượn tài sản",
    "mượn": "Thuê - Mượn tài sản",
}


CRITICAL_WEB_FORM_FIELDS = (
    "ten_hop_dong",
    "so_cong_chung",
    "nhom_hop_dong",
    "loai_tai_san",
    "tai_san",
)

DOC_KIND_TRANSFER = "transfer_contract"
DOC_KIND_ASSET_COMMITMENT = "asset_commitment"
DOC_KIND_TRANSFER_CANCELLATION = "transfer_cancellation"
DOC_KIND_GENERIC = "generic"


def _append_text_lines(lines: list[str], values: Iterable[str]) -> None:
    for raw in values:
        text = str(raw or "").strip()
        if text:
            lines.append(text)


def _append_table_lines(lines: list[str], tables) -> None:
    for table in tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text and p.text.strip()
                ).strip()
                if cell_text:
                    row_values.append(cell_text)
            if row_values:
                lines.append(" | ".join(row_values))


# ============================================================
# DOCX / DOC HELPERS
# ============================================================
def read_doc_via_ifilter(filepath) -> str:
    """Trich xuat text tu .doc bang Windows IFilter (cung co che voi Agent Ransack).
    WHY: Nhanh hon Word COM 10-20x (0.1-0.3s/file vs 3-5s/file), khong can mo Word UI.
         Dung cho scan (tim so hop dong) — chi can plain text, khong can table structure.
    RISK: IFilter tra plain text thuan — mat cau truc bang va header/footer.
          Khong dung cho extract() (can table), chi dung cho scan_docx_for_contract_no().
    REQUIRE: query.dll (co san tren Windows 7+, khong can cai them).
    """
    import ctypes

    abs_path = str(Path(filepath).resolve())
    ole32 = ctypes.windll.ole32
    ole32.CoInitialize(None)
    try:
        query = ctypes.windll.query
        # c_long (signed) de compare HRESULT am — WINFUNCTYPE(HRESULT) tu raise exception
        query.LoadIFilter.restype = ctypes.c_long
        query.LoadIFilter.argtypes = [
            ctypes.c_wchar_p,
            ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_void_p),
        ]
        p_filter = ctypes.c_void_p()
        hr = query.LoadIFilter(abs_path, None, ctypes.byref(p_filter))
        if hr != 0 or not p_filter.value:
            raise OSError(f"LoadIFilter that bai (hr=0x{hr & 0xFFFFFFFF:08X})")

        # IFilter VTBL layout (IUnknown 0-2, IFilter 3-7):
        #   3=Init, 4=GetChunk, 5=GetText, 6=GetValue, 7=BindRegion
        VTBL_INIT = 3
        VTBL_GET_CHUNK = 4
        VTBL_GET_TEXT = 5
        VTBL_RELEASE = 2

        vtbl_ptr = ctypes.cast(p_filter, ctypes.POINTER(ctypes.c_void_p))
        vtbl = ctypes.cast(vtbl_ptr[0], ctypes.POINTER(ctypes.c_void_p))

        # Init(grfFlags, cAttributes, aAttributes, pFlags) — c_long tranh auto-raise
        InitFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p,
            ctypes.c_uint, ctypes.c_ulong, ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_uint),
        )
        out_flags = ctypes.c_uint(0)
        InitFunc(vtbl[VTBL_INIT])(p_filter, 0, 0, None, ctypes.byref(out_flags))

        # STAT_CHUNK tren x64 Windows:
        #   offset  0: idChunk   (ULONG 4B)
        #   offset  4: breakType (enum  4B)
        #   offset  8: flags     (CHUNKSTATE 4B) — CHUNK_TEXT=0x1, CHUNK_VALUE=0x2
        #   offset 12: locale    (LCID  4B)
        #   offset 16: attribute (FULLPROPSPEC = GUID(16B) + PROPSPEC(4+4pad+8B) = 32B)
        #   offset 48: idChunkSource (4B), cwcStartSource (4B), cwcLenSource (4B)
        #   total: 60B — dung 64B buffer cho an toan tren moi alignment
        CHUNK_BUF_SIZE = 64
        CHUNK_FLAGS_OFFSET = 8
        CHUNK_TEXT = 0x1  # CHUNKSTATE: CHUNK_TEXT=1 (khong phai 2)

        GetChunkFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p, ctypes.c_char_p,
        )
        get_chunk_fn = GetChunkFunc(vtbl[VTBL_GET_CHUNK])

        BUF_SIZE = 4096
        GetTextFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_ulong), ctypes.c_wchar_p,
        )
        get_text_fn = GetTextFunc(vtbl[VTBL_GET_TEXT])

        ReleaseFunc = ctypes.WINFUNCTYPE(ctypes.c_ulong, ctypes.c_void_p)
        release_fn = ReleaseFunc(vtbl[VTBL_RELEASE])

        # HRESULT codes (signed 32-bit)
        FILTER_E_END_OF_CHUNKS = ctypes.c_long(0x80041700).value  # negative
        FILTER_E_NO_MORE_TEXT  = ctypes.c_long(0x80041701).value  # negative
        # FILTER_S_LAST_TEXT = 0x00041709 — SUCCESS: text duoc ghi VA day la phan cuoi cung
        # Phai capture text TRUOC khi check status nay, roi break
        FILTER_S_LAST_TEXT = 0x00041709  # positive (success with info)

        parts: list[str] = []
        chunk_buf = ctypes.create_string_buffer(CHUNK_BUF_SIZE)

        while True:
            hr_chunk = get_chunk_fn(p_filter, chunk_buf)
            if hr_chunk == FILTER_E_END_OF_CHUNKS:
                break
            if hr_chunk != 0:
                break

            flags = ctypes.c_uint.from_buffer_copy(chunk_buf, CHUNK_FLAGS_OFFSET).value
            if not (flags & CHUNK_TEXT):
                continue

            while True:
                buf_len = ctypes.c_ulong(BUF_SIZE)
                text_buf = ctypes.create_unicode_buffer(BUF_SIZE)
                hr_text = get_text_fn(p_filter, ctypes.byref(buf_len), text_buf)
                # Capture text TRUOC khi check status — FILTER_S_LAST_TEXT tra text hop le
                if buf_len.value > 0:
                    parts.append(text_buf.value[: buf_len.value])
                if hr_text == FILTER_E_NO_MORE_TEXT:
                    break
                if hr_text == FILTER_S_LAST_TEXT:
                    break  # text da duoc lay, day la phan cuoi cua chunk nay
                if hr_text != 0:
                    break

        release_fn(p_filter)
        return "".join(parts)
    finally:
        ole32.CoUninitialize()


def read_doc_via_word_com(filepath) -> str:
    """Doc .doc bang Word COM — lay ca table + header/footer.
    WHY: Dung cho extract() can cau truc day du. Cham (3-5s/file) nhung chinh xac nhat.
    RISK: neu Word dang bi lock file thi se raise; caller nen catch Exception.
    """
    import tempfile
    import pythoncom
    import win32com.client as win32

    abs_path = str(Path(filepath).resolve())
    pythoncom.CoInitialize()
    word = None
    doc = None
    tmp_path = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_path, ReadOnly=True, AddToRecentFiles=False)

        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(tmp_fd)
        # wdFormatXMLDocument = 12
        doc.SaveAs2(tmp_path, FileFormat=12)
        doc.Close(SaveChanges=False)
        doc = None

        from docx import Document
        d = Document(tmp_path)
        lines: list[str] = []
        _append_text_lines(lines, (p.text for p in d.paragraphs))
        _append_table_lines(lines, d.tables)
        for section in d.sections:
            _append_text_lines(lines, (p.text for p in section.header.paragraphs))
            _append_text_lines(lines, (p.text for p in section.footer.paragraphs))
        return "\n".join(lines)
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if tmp_path and Path(tmp_path).exists():
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def read_docx(filepath, *, use_ifilter_for_doc: bool = False):
    """Doc .docx hoac .doc -> text thuan.
    - .docx: dung python-docx truc tiep (nhanh).
    - .doc + use_ifilter_for_doc=True: IFilter (nhanh, plain text).
    - .doc + use_ifilter_for_doc=False: Word COM (cham, day du cau truc).
    """
    path = Path(filepath)
    if path.suffix.lower() == ".doc":
        if use_ifilter_for_doc:
            return read_doc_via_ifilter(filepath)
        return read_doc_via_word_com(filepath)

    from docx import Document
    doc = Document(str(filepath))
    lines: list[str] = []
    _append_text_lines(lines, (p.text for p in doc.paragraphs))
    _append_table_lines(lines, doc.tables)
    for section in doc.sections:
        _append_text_lines(lines, (p.text for p in section.header.paragraphs))
        _append_text_lines(lines, (p.text for p in section.footer.paragraphs))
    return "\n".join(lines)


def batch_convert_doc_to_docx(
    folder: Path,
    *,
    skip_existing: bool = True,
    log_callback=None,
) -> dict:
    """Convert toan bo .doc trong folder sang .docx cung vi tri, dung 1 Word instance.
    WHY: Sau khi convert 1 lan, moi lan scan sau doc .docx truc tiep (khong can Word/IFilter).
    File .doc goc giu nguyen. Bo qua ~$ temp file.
    Returns: {"converted": int, "skipped": int, "failed": int, "errors": list}
    """
    import pythoncom
    import win32com.client as win32

    def log(msg: str) -> None:
        if log_callback:
            log_callback(msg)

    folder = Path(folder)
    doc_files = sorted(
        f for f in folder.rglob("*.doc")
        if not f.name.startswith("~$")
    )
    total = len(doc_files)
    if total == 0:
        log("[CONVERT] Khong tim thay file .doc nao.")
        return {"converted": 0, "skipped": 0, "failed": 0, "errors": []}

    log(f"[CONVERT] Tim thay {total} file .doc. Bat dau chuyen doi...")

    stats = {"converted": 0, "skipped": 0, "failed": 0, "errors": []}

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        for idx, doc_path in enumerate(doc_files, start=1):
            prefix = f"[CONVERT] {idx}/{total}"
            docx_path = doc_path.with_suffix(".docx")
            doc = None

            if skip_existing and docx_path.exists():
                log(f"{prefix} skip (da co docx): {doc_path.name}")
                stats["skipped"] += 1
                continue

            try:
                doc = word.Documents.Open(
                    str(doc_path.resolve()),
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                # wdFormatXMLDocument = 12
                doc.SaveAs2(str(docx_path.resolve()), FileFormat=12)
                doc.Close(SaveChanges=False)
                doc = None
                log(f"{prefix} OK: {doc_path.name} -> {docx_path.name}")
                stats["converted"] += 1
            except Exception as exc:
                log(f"{prefix} LOI: {doc_path.name} — {exc}")
                stats["errors"].append({"file": str(doc_path), "error": str(exc)})
                stats["failed"] += 1
            finally:
                if doc is not None:
                    try:
                        doc.Close(SaveChanges=False)
                    except Exception:
                        pass
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    log(
        f"[CONVERT] Xong. converted={stats['converted']}, "
        f"skipped={stats['skipped']}, failed={stats['failed']}"
    )
    return stats


def _normalize_scan_contract_no(contract_no: str) -> str:
    return str(contract_no or "").strip().upper()


def _normalize_web_contract_no(contract_no: str) -> str:
    normalized = _normalize_scan_contract_no(contract_no)
    if normalized.endswith("/CCGD"):
        normalized = normalized[:-5]
    return normalized


def find_contract_numbers(text: str) -> list[str]:
    return [_normalize_scan_contract_no(match.group(1)) for match in CONTRACT_NO_REGEX.finditer(text or "")]


def has_contract_keyword(text: str, file_name: str = "") -> bool:
    haystack = f"{file_name}\n{text}".lower()
    return any(keyword in haystack for keyword in CONTRACT_KEYWORDS)


def _normalize_plain_text_for_extract(text: str) -> str:
    normalized = str(text or "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\x0b", "\n").replace("\t", " ").replace("\xa0", " ")
    lines = []
    for raw in normalized.split("\n"):
        line = re.sub(r"\s+", " ", raw).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _scan_contract_text(text: str, *, file_name: str = "") -> dict:
    matches = find_contract_numbers(text)
    if matches:
        contract_no = matches[-1]
        return {
            "is_contract": True,
            "contract_no": contract_no,
            "reason": f"Tim thay so CC: {contract_no}",
        }

    if has_contract_keyword(text, file_name):
        return {
            "is_contract": False,
            "contract_no": "",
            "reason": "Co keyword hop dong nhung khong thay so cong chung",
        }

    return {
        "is_contract": False,
        "contract_no": "",
        "reason": "Khong thay so cong chung",
    }


def scan_docx_for_contract_no(filepath, *, include_text: bool = False):
    """Tim so cong chung trong noi dung file .docx hoac .doc.
    .doc dung IFilter (nhanh) — chi can plain text de tim regex so cong chung.
    """
    path = Path(filepath)
    use_ifilter_for_doc = path.suffix.lower() == ".doc"
    try:
        text = read_docx(path, use_ifilter_for_doc=use_ifilter_for_doc)
        if use_ifilter_for_doc:
            text = _normalize_plain_text_for_extract(text)
    except Exception as exc:
        result = {
            "is_contract": False,
            "contract_no": "",
            "reason": f"Khong doc duoc file: {exc}",
        }
        if include_text:
            result["text"] = ""
        return result

    result = _scan_contract_text(text, file_name=path.name)
    if include_text:
        result["text"] = text
    return result


# ============================================================
# TRICH XUAT TUNG TRUONG
# ============================================================
def find_ten_hop_dong(text):
    return _detect_document_kind_and_title(text)[1]


def find_so_cong_chung(text):
    m = re.search(r"Số công chứng\s*(\d+)\s*/\s*(\d{4})\s*/\s*([A-Za-z]+)", text)
    if m:
        return _normalize_web_contract_no(f"{m.group(1)}/{m.group(2)}/{m.group(3)}")

    idx = text.find("Số công chứng")
    if idx >= 0:
        fallback = re.search(r"(\d+)/(\d{4})", text[idx: idx + 100])
        if fallback:
            return _normalize_web_contract_no(f"{fallback.group(1)}/{fallback.group(2)}")

    matches = find_contract_numbers(text)
    if matches:
        return _normalize_web_contract_no(matches[-1])
    return ""


def find_ngay_cong_chung(text):
    idx = text.find("LỜI CHỨNG")
    sub = text[idx:] if idx >= 0 else text
    m = re.search(r"ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})", sub)
    if m:
        return f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/{m.group(3)}"
    return ""


def find_ccv(text):
    m = re.search(r"Tôi,\s+([^,]+),\s+Công chứng viên", text)
    return m.group(1).strip() if m else DEFAULT_CCV


def find_persons(section):
    """Tim tat ca nguoi trong 1 doan (Ben A hoac Ben B)."""
    persons = []
    names = list(re.finditer(r"(Ông|[Bb]à)\s+([^;]+?)\s*;", section))

    for idx, nm in enumerate(names):
        start = nm.start()
        end = names[idx + 1].start() if idx + 1 < len(names) else len(section)
        chunk = section[start:end]

        person = {
            "gioi_tinh": nm.group(1).capitalize(),
            "ho_ten": nm.group(2).strip(),
        }

        m = re.search(r"Sinh ngày:?\s*(\d{2}/\d{2}/\d{4})", chunk)
        person["ngay_sinh"] = m.group(1) if m else ""

        m = re.search(r"(?:Căn cước|CCCD|CMND)(?:\s+công dân)?\s*(?:số:?\s*)(\d+)", chunk)
        person["cccd"] = m.group(1) if m else ""

        m = re.search(r"do\s+(.+?)\s+cấp ngày", chunk, re.DOTALL)
        person["noi_cap"] = re.sub(r"\s+", " ", m.group(1)).strip() if m else ""

        m = re.search(r"cấp ngày\s*(\d{2}/\d{2}/\d{4})", chunk)
        person["ngay_cap_cccd"] = m.group(1) if m else ""

        persons.append(person)

    return persons


def find_dia_chi(section):
    m = re.search(r"cùng (?:cư trú|thường trú) tại:?\s*(.+?)\.", section, re.DOTALL)
    return re.sub(r"\s+", " ", m.group(1)).strip() if m else ""


def find_tai_san(text):
    doc_kind, _ = _detect_document_kind_and_title(text)
    return _find_tai_san_by_kind(text, doc_kind)


def _fold_text(text: str) -> str:
    normalized = unicodedata.normalize("NFD", str(text or ""))
    normalized = normalized.replace("\u0111", "d").replace("\u0110", "D")
    return "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn").lower()


def _sentence_case_vn(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip(" .:-;\n\t")
    if not cleaned:
        return ""
    return cleaned[:1].upper() + cleaned[1:].lower()


def _find_title_line(text: str) -> str:
    for raw in str(text or "").split("\n")[:20]:
        line = re.sub(r"\s+", " ", raw).strip(" .:-;")
        if not line:
            continue
        folded = _fold_text(line)
        if (
            folded.startswith("hop dong ")
            or folded.startswith("van ban ")
            or folded.startswith("thoa thuan ")
            or "cam ket" in folded
        ):
            return line
    return ""


def _detect_document_kind_and_title(text: str, *, file_name: str = "") -> tuple[str, str]:
    title_line = _find_title_line(text)
    title_fold = _fold_text(title_line)
    combined_fold = _fold_text(f"{file_name}\n{text}")

    if "cam ket tai san rieng" in title_fold or "cam ket tai san rieng" in combined_fold:
        return DOC_KIND_ASSET_COMMITMENT, "Văn bản cam kết tài sản riêng"

    is_transfer_cancellation = (
        ("huy bo" in title_fold or "huy bo" in combined_fold or "huy hop dong" in combined_fold)
        and "chuyen nhuong" in combined_fold
    )
    if is_transfer_cancellation:
        if any(
            marker in combined_fold
            for marker in (
                "tai san gan lien voi dat",
                "nha o va quyen su dung dat o",
                "mua ban nha o",
            )
        ):
            return (
                DOC_KIND_TRANSFER_CANCELLATION,
                "Văn bản thỏa thuận về việc hủy bỏ hợp đồng chuyển nhượng quyền sử dụng đất và tài sản gắn liền với đất",
            )
        return (
            DOC_KIND_TRANSFER_CANCELLATION,
            "Văn bản thỏa thuận về việc hủy bỏ hợp đồng chuyển nhượng quyền sử dụng đất",
        )

    if "hop dong chuyen nhuong" in title_fold or "hop dong chuyen nhuong" in combined_fold:
        if any(
            marker in combined_fold
            for marker in (
                "tai san gan lien voi dat",
                "nha o va quyen su dung dat o",
                "va nha o",
                "va tai san",
            )
        ):
            return DOC_KIND_TRANSFER, "Hợp đồng chuyển nhượng quyền sử dụng đất và tài sản gắn liền với đất"
        return DOC_KIND_TRANSFER, "Hợp đồng chuyển nhượng quyền sử dụng đất"

    if title_line:
        return DOC_KIND_GENERIC, _sentence_case_vn(title_line)
    return DOC_KIND_GENERIC, ""


def _extract_block_by_patterns(text: str, start_patterns: Iterable[str], end_patterns: Iterable[str]) -> str:
    match = _find_first_pattern_match_generic(text, start_patterns)
    if not match:
        return ""

    start_idx = int(match[0])
    end_idx = len(text)
    for pattern in end_patterns:
        end_match = re.search(pattern, text[start_idx:], re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not end_match:
            continue
        candidate_end = start_idx + end_match.start()
        if candidate_end > start_idx and candidate_end < end_idx:
            end_idx = candidate_end

    raw = text[start_idx:end_idx].strip()
    if not raw:
        return ""
    lines = [re.sub(r"\s+", " ", line).strip(" ;,") for line in raw.split("\n") if line.strip()]
    return "\n".join(lines).strip()


def _find_tai_san_generic(text: str) -> str:
    return _extract_block_by_patterns(
        text,
        (
            r"(?m)^\s*(?:1\.1\s*)?\u0110\u1ed1i t\u01b0\u1ee3ng c\u1ee7a H\u1ee3p \u0111\u1ed3ng n\u00e0y l\u00e0\b",
            r"(?m)^\s*\u0110I\u1ec0U\s*1\b",
        ),
        (
            r"(?m)^\s*1\.2\b",
            r"(?m)^\s*B\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
            r"(?m)^\s*\u0110I\u1ec0U\s*2\b",
            r"(?m)^\s*Gi\u00e1 chuy\u1ec3n nh\u01b0\u1ee3ng\b",
            r"(?m)^\s*Ph\u01b0\u01a1ng th\u1ee9c thanh to\u00e1n\b",
        ),
    )


def _find_tai_san_by_kind(text: str, doc_kind: str) -> str:
    if doc_kind == DOC_KIND_TRANSFER:
        return _extract_block_by_patterns(
            text,
            (r"(?m)^\s*(?:1\.1\s*)?\u0110\u1ed1i t\u01b0\u1ee3ng c\u1ee7a H\u1ee3p \u0111\u1ed3ng n\u00e0y l\u00e0\b",),
            (
                r"(?m)^\s*1\.2\b",
                r"(?m)^\s*B\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
                r"(?m)^\s*\u0110I\u1ec0U\s*2\b",
                r"(?m)^\s*Gi\u00e1 chuy\u1ec3n nh\u01b0\u1ee3ng\b",
                r"(?m)^\s*Ph\u01b0\u01a1ng th\u1ee9c thanh to\u00e1n\b",
            ),
        )

    if doc_kind == DOC_KIND_ASSET_COMMITMENT:
        return _extract_block_by_patterns(
            text,
            (
                r"(?m)^\s*(?:\u00d4ng|B\u00e0)[^\n]{0,200}?hi\u1ec7n \u0111ang s\u1edf h\u1eefu\s+T\u00e0i\s*S\u1ea3n\s+l\u00e0\b",
                r"(?m)^\s*T\u00e0i\s*S\u1ea3n\s+l\u00e0\b",
            ),
            (
                r"(?m)^\s*B\u1eb1ng v\u0103n b\u1ea3n n\u00e0y ch\u00fang t\u00f4i x\u00e1c \u0111\u1ecbnh\b",
                r"(?m)^\s*Hai v\u1ee3 ch\u1ed3ng ch\u00fang t\u00f4i cam \u0111oan\b",
                r"(?m)^\s*Ch\u00fang t\u00f4i c\u00f4ng nh\u1eadn\b",
            ),
        )

    if doc_kind == DOC_KIND_TRANSFER_CANCELLATION:
        return _extract_block_by_patterns(
            text,
            (
                r"(?:mua b\u00e1n nh\u00e0 \u1edf v\u00e0\s+)?chuy\u1ec3n nh\u01b0\u1ee3ng quy\u1ec1n s\u1eed d\u1ee5ng \u0111\u1ea5t(?:\s+v\u00e0\s+t\u00e0i\s+s\u1ea3n\s+g\u1eafn\s+li\u1ec1n\s+v\u1edbi\s+\u0111\u1ea5t)?[^.\n]{0,200}?c\u00f3 \u0111\u1ecba ch\u1ec9 t\u1ea1i:\s*",
            ),
            (
                r"\bv\u00e0 \u0111\u01b0\u1ee3c C\u00f4ng ch\u1ee9ng vi\u00ean\b",
                r"\bv\u00e0 \u0111\u01b0\u1ee3c[^.\n]{0,120}?ch\u1ee9ng nh\u1eadn\b",
                r"\bs\u1ed1 c\u00f4ng ch\u1ee9ng\b",
            ),
        )

    return _find_tai_san_generic(text)


def guess_nhom_hd(ten_hd):
    lower = ten_hd.lower()
    for kw, nhom in NHOM_HD_MAP.items():
        if kw in lower:
            return nhom
    return "Khác"


def guess_loai_tai_san(tai_san, ten_hd):
    combined = _fold_text(f"{tai_san} {ten_hd}")

    if "quyen su dung dat" in combined:
        indicators = (
            "va nha o",
            "va tai san",
            "nha o gan lien",
            "cong trinh xay dung",
            "tai san tren dat",
            "tai san gan lien voi dat",
            "nha o va quyen su dung dat o",
        )
        if any(keyword in combined for keyword in indicators):
            return "Đất đai có tài sản"
        return "Đất đai không có tài sản"

    keywords = [
        ("o to", "Ô tô"),
        ("xe may", "Xe máy"),
        ("tau", "Tàu, thuyền"),
        ("thuyen", "Tàu, thuyền"),
        ("so tiet kiem", "Sổ tiết kiệm"),
        ("co phan", "Cổ phần"),
        ("co phieu", "Cổ phiếu"),
        ("chung khoan", "Chứng khoán"),
        ("the atm", "Thẻ ATM"),
    ]
    for kw, loai in keywords:
        if kw in combined:
            return loai
    return "Tài sản khác"


# ============================================================
# FORMAT CHO 3 O CKEDITOR TREN WEB
# ============================================================
def fmt_nguoi_yeu_cau(ben_b):
    """Nguoi yeu cau CC = nguoi dau tien ben B."""
    if not ben_b.get("nguoi"):
        return ""
    person = ben_b["nguoi"][0]
    parts = [f"{person.get('gioi_tinh', '')} {person['ho_ten']}"]
    if person.get("ngay_sinh"):
        parts.append(f"Sinh ngày: {person['ngay_sinh']}")
    if person.get("cccd"):
        parts.append(f"Căn cước số: {person['cccd']}")
    if person.get("noi_cap") and person.get("ngay_cap_cccd"):
        parts.append(f"do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}")
    if ben_b.get("dia_chi"):
        parts.append(f"Cư trú tại: {ben_b['dia_chi']}")
    return "; ".join(parts)


def fmt_duong_su(ben_a, ben_b):
    """Duong su = toan bo ben A + ben B."""
    lines = ["BÊN A (Bên chuyển nhượng):"]
    for idx, person in enumerate(ben_a.get("nguoi", []), 1):
        line = f"{idx}. {person.get('gioi_tinh', '')} {person['ho_ten']}"
        if person.get("ngay_sinh"):
            line += f"; Sinh ngày: {person['ngay_sinh']}"
        if person.get("cccd"):
            line += f"; CCCD: {person['cccd']}"
        if person.get("noi_cap") and person.get("ngay_cap_cccd"):
            line += f"; do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}"
        lines.append(line)
    if ben_a.get("dia_chi"):
        lines.append(f"Cùng cư trú tại: {ben_a['dia_chi']}")

    lines.append("")
    lines.append("BÊN B (Bên nhận chuyển nhượng):")
    for idx, person in enumerate(ben_b.get("nguoi", []), 1):
        line = f"{idx}. {person.get('gioi_tinh', '')} {person['ho_ten']}"
        if person.get("ngay_sinh"):
            line += f"; Sinh ngày: {person['ngay_sinh']}"
        if person.get("cccd"):
            line += f"; CCCD: {person['cccd']}"
        if person.get("noi_cap") and person.get("ngay_cap_cccd"):
            line += f"; do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}"
        lines.append(line)
    if ben_b.get("dia_chi"):
        lines.append(f"Cùng cư trú tại: {ben_b['dia_chi']}")

    return "\n".join(lines)


def get_missing_web_form_fields(web_form: dict, *, file_hop_dong: str = "") -> list[str]:
    values = dict(web_form or {})
    values["file_hop_dong"] = file_hop_dong
    return [field for field in CRITICAL_WEB_FORM_FIELDS if not str(values.get(field) or "").strip()]


def _find_first_pattern_index(text: str, patterns: Iterable[str], *, start: int = 0) -> int:
    haystack = text[start:]
    best_idx = -1
    for pattern in patterns:
        match = re.search(pattern, haystack, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not match:
            continue
        idx = start + match.start()
        if best_idx < 0 or idx < best_idx:
            best_idx = idx
    return best_idx


def _extract_plain_text_persons(section: str) -> list[dict]:
    header_re = re.compile(
        r"(?i)^\s*(?:\d+\.\s*)?(?:v\u00e0\s+(?:v\u1ee3|ch\u1ed3ng)\s+l\u00e0\s+)?"
        r"(?:ng\u01b0\u1eddi\s+(?:v\u1ee3|ch\u1ed3ng)\s+)?(?P<title>\u00f4ng|b\u00e0)\s*:?\s*"
        r"(?P<name>.+?)(?=\s+Sinh ng\u00e0y:|[;]|$)"
    )
    birth_re = re.compile(r"(?i)Sinh ng\u00e0y:?\s*(\d{1,2}/\d{1,2}/\d{4})")
    id_re = re.compile(r"(?i)(?:C\u0103n c\u01b0\u1edbc(?:\s+c\u00f4ng d\u00e2n)?|CCCD|CMND)\s*(?:s\u1ed1)?\s*:?\s*(\d+)")
    issue_place_re = re.compile(r"(?i)\bdo\s+(.+?)\s+c\u1ea5p ng\u00e0y", re.DOTALL)
    issue_date_re = re.compile(r"(?i)c\u1ea5p ng\u00e0y\s*(\d{1,2}/\d{1,2}/\d{4})")

    lines = [line.strip() for line in section.split("\n") if line.strip()]
    headers: list[tuple[int, re.Match[str]]] = []
    for idx, line in enumerate(lines):
        match = header_re.search(line)
        if match:
            headers.append((idx, match))

    persons = []
    for header_idx, (line_idx, match) in enumerate(headers):
        next_idx = headers[header_idx + 1][0] if header_idx + 1 < len(headers) else len(lines)
        chunk = "\n".join(lines[line_idx:next_idx])
        title = match.group("title").strip().lower()
        person = {
            "gioi_tinh": "\u00d4ng" if title.startswith("\u00f4") else "B\u00e0",
            "ho_ten": re.sub(r"\s+", " ", match.group("name")).strip(" ;,."),
        }
        birth_match = birth_re.search(chunk)
        person["ngay_sinh"] = birth_match.group(1) if birth_match else ""
        id_match = id_re.search(chunk)
        person["cccd"] = id_match.group(1) if id_match else ""
        issue_place_match = issue_place_re.search(chunk)
        person["noi_cap"] = re.sub(r"\s+", " ", issue_place_match.group(1)).strip(" ;,.") if issue_place_match else ""
        issue_date_match = issue_date_re.search(chunk)
        person["ngay_cap_cccd"] = issue_date_match.group(1) if issue_date_match else ""
        name_fold = _fold_text(person["ho_ten"])
        if "hien dang so huu" in name_fold or "tai san la" in name_fold:
            continue
        if not birth_match and not id_match:
            continue
        persons.append(person)
    return persons


def _extract_plain_text_address(section: str) -> str:
    match = re.search(
        r"(?im)^(?:C\u1ea3 hai \u00f4ng b\u00e0\s+)?(?:C\u00f9ng\s+)?(?:N\u01a1i\s+)?"
        r"(?:th\u01b0\u1eddng tr\u00fa|c\u01b0 tr\u00fa)\s+t\u1ea1i:\s*(.+)$",
        section,
    )
    if not match:
        return ""
    return re.sub(r"\s+", " ", match.group(1)).strip(" .;")


def _extract_plain_text_parties(text: str) -> tuple[dict, dict]:
    ben_a_patterns = (r"\b(?:I\.\s*)?B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG\b", r"\bB\u00caN A\b")
    ben_b_patterns = (r"\b(?:II\.\s*)?B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG\b", r"\bB\u00caN B\b")
    end_patterns = (
        r"\bHai b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bB\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
        r"\b\u0110I\u1ec0U\s*1\b",
        r"\bL\u1edcI CH\u1ee8NG\b",
    )

    idx_a = _find_first_pattern_index(text, ben_a_patterns)
    idx_b = _find_first_pattern_index(text, ben_b_patterns, start=max(idx_a, 0))
    idx_end = _find_first_pattern_index(text, end_patterns, start=max(idx_b, 0)) if idx_b >= 0 else -1

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end > idx_b else text[idx_b:] if idx_b >= 0 else ""

    ben_a = {
        "nguoi": _extract_plain_text_persons(ben_a_section),
        "dia_chi": _extract_plain_text_address(ben_a_section),
    }
    ben_b = {
        "nguoi": _extract_plain_text_persons(ben_b_section),
        "dia_chi": _extract_plain_text_address(ben_b_section),
    }
    return ben_a, ben_b


def _find_first_pattern_match_generic(text: str, patterns: Iterable[str], *, start: int = 0):
    haystack = text[start:]
    best_match = None
    for pattern in patterns:
        match = re.search(pattern, haystack, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not match:
            continue
        idx = start + match.start()
        end = start + match.end()
        if best_match is None or idx < best_match[0]:
            best_match = (idx, end)
    return best_match


def _extract_parties_generic(text: str) -> tuple[dict, dict]:
    ben_a_patterns = (
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHO VAY\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHO THU\u00ca\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN U[\u1ef6Y\u1ee6] QUY\u1ec0N\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN T\u1eb6NG CHO\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN B\u00c1N\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN A\b",
    )
    ben_b_patterns = (
        r"(?m)^\s*(?:II\.\s*)?B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN VAY\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN THU\u00ca\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN \u0110\u01af\u1ee2C U[\u1ef6Y\u1ee6] QUY\u1ec0N\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN NH\u1eacN T\u1eb6NG CHO\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN MUA\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN B\b",
    )
    preamble_patterns = (
        r"\bCh\u00fang t\u00f4i g\u1ed3m c\u00f3\s*:\s*",
        r"\bCh\u00fang t\u00f4i g\u1ed3m\s*:\s*",
    )
    end_patterns = (
        r"\bC\u00e1c b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bHai b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bB\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
        r"\b\u0110I\u1ec0U\s*1\b",
        r"\bL\u1edcI CH\u1ee8NG\b",
    )

    match_a = _find_first_pattern_match_generic(text, ben_a_patterns)
    match_b = _find_first_pattern_match_generic(text, ben_b_patterns)
    preamble_match = _find_first_pattern_match_generic(text, preamble_patterns)

    if match_a and match_b and match_a[0] > match_b[0]:
        match_a = None

    if match_a:
        idx_a_start = int(match_a[0])
    elif preamble_match and match_b:
        idx_a_start = int(preamble_match[1])
    else:
        idx_a_start = 0

    idx_b_start = int(match_b[0]) if match_b else -1
    idx_end = _find_first_pattern_index(text, end_patterns, start=max(idx_b_start, idx_a_start))

    if idx_b_start >= 0:
        ben_a_section = text[idx_a_start:idx_b_start]
        ben_b_section = text[idx_b_start:idx_end] if idx_end > idx_b_start else text[idx_b_start:]
    else:
        ben_a_section = text[idx_a_start:idx_end] if idx_end > idx_a_start else text[idx_a_start:]
        ben_b_section = ""

    ben_a = {
        "nguoi": _extract_plain_text_persons(ben_a_section),
        "dia_chi": _extract_plain_text_address(ben_a_section),
    }
    ben_b = {
        "nguoi": _extract_plain_text_persons(ben_b_section),
        "dia_chi": _extract_plain_text_address(ben_b_section),
    }
    return ben_a, ben_b


def _fmt_duong_su_generic(ben_a, ben_b):
    lines = ["BÊN A:"]
    for idx, person in enumerate(ben_a.get("nguoi", []), 1):
        line = f"{idx}. {person.get('gioi_tinh', '')} {person['ho_ten']}"
        if person.get("ngay_sinh"):
            line += f"; Sinh ngày: {person['ngay_sinh']}"
        if person.get("cccd"):
            line += f"; CCCD: {person['cccd']}"
        if person.get("noi_cap") and person.get("ngay_cap_cccd"):
            line += f"; do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}"
        lines.append(line)
    if ben_a.get("dia_chi"):
        lines.append(f"Cùng cư trú tại: {ben_a['dia_chi']}")

    lines.append("")
    lines.append("BÊN B:")
    for idx, person in enumerate(ben_b.get("nguoi", []), 1):
        line = f"{idx}. {person.get('gioi_tinh', '')} {person['ho_ten']}"
        if person.get("ngay_sinh"):
            line += f"; Sinh ngày: {person['ngay_sinh']}"
        if person.get("cccd"):
            line += f"; CCCD: {person['cccd']}"
        if person.get("noi_cap") and person.get("ngay_cap_cccd"):
            line += f"; do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}"
        lines.append(line)
    if ben_b.get("dia_chi"):
        lines.append(f"Cùng cư trú tại: {ben_b['dia_chi']}")
    return "\n".join(lines)


def _build_payload_generic(filepath, text: str, scan_result: dict, *, extract_mode: str) -> dict:
    ben_a, ben_b = _extract_parties_generic(text)
    doc_kind, ten_hd = _detect_document_kind_and_title(text, file_name=Path(filepath).name)
    tai_san = _find_tai_san_by_kind(text, doc_kind)
    so_cong_chung = _normalize_web_contract_no(find_so_cong_chung(text) or scan_result.get("contract_no", ""))
    nguoi_yeu_cau_party = ben_b if ben_b.get("nguoi") else ben_a
    web_form = {
        "ten_hop_dong": ten_hd,
        "ngay_cong_chung": find_ngay_cong_chung(text),
        "so_cong_chung": so_cong_chung,
        "nhom_hop_dong": guess_nhom_hd(ten_hd),
        "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
        "cong_chung_vien": find_ccv(text),
        "thu_ky": DEFAULT_THU_KY,
        "nguoi_yeu_cau": fmt_nguoi_yeu_cau(nguoi_yeu_cau_party),
        "duong_su": _fmt_duong_su_generic(ben_a, ben_b),
        "tai_san": tai_san,
    }
    file_goc = os.path.abspath(filepath)
    missing_fields = get_missing_web_form_fields(web_form, file_hop_dong=file_goc)
    return {
        "web_form": web_form,
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": file_goc,
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
            "document_kind": doc_kind,
            "extract_mode": extract_mode,
            "extract_is_partial": bool(missing_fields),
            "missing_web_form_fields": missing_fields,
        },
    }


def _build_payload(filepath, text: str, scan_result: dict, ben_a: dict, ben_b: dict, *, extract_mode: str) -> dict:
    doc_kind, ten_hd = _detect_document_kind_and_title(text, file_name=Path(filepath).name)
    tai_san = _find_tai_san_by_kind(text, doc_kind)
    so_cong_chung = _normalize_web_contract_no(find_so_cong_chung(text) or scan_result.get("contract_no", ""))
    nguoi_yeu_cau_party = ben_b if ben_b.get("nguoi") else ben_a
    web_form = {
        "ten_hop_dong": ten_hd,
        "ngay_cong_chung": find_ngay_cong_chung(text),
        "so_cong_chung": so_cong_chung,
        "nhom_hop_dong": guess_nhom_hd(ten_hd),
        "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
        "cong_chung_vien": find_ccv(text),
        "thu_ky": DEFAULT_THU_KY,
        "nguoi_yeu_cau": fmt_nguoi_yeu_cau(nguoi_yeu_cau_party),
        "duong_su": fmt_duong_su(ben_a, ben_b),
        "tai_san": tai_san,
    }
    file_goc = os.path.abspath(filepath)
    missing_fields = get_missing_web_form_fields(web_form, file_hop_dong=file_goc)
    return {
        "web_form": web_form,
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": file_goc,
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
            "document_kind": doc_kind,
            "extract_mode": extract_mode,
            "extract_is_partial": bool(missing_fields),
            "missing_web_form_fields": missing_fields,
        },
    }


def _extract_structured_text_payload(filepath, text: str, scan_result: dict) -> dict:
    idx_a = text.find("BÃŠN CHUYá»‚N NHÆ¯á»¢NG")
    idx_b = text.find("BÃŠN NHáº¬N CHUYá»‚N NHÆ¯á»¢NG")
    idx_end = text.find("Hai bÃªn tá»± nguyá»‡n")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="structured_docx")


def _extract_plain_text_payload(filepath, text: str, scan_result: dict) -> dict:
    ben_a, ben_b = _extract_plain_text_parties(text)
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="plain_text_doc")


def _extract_structured_text_payload_v2(filepath, text: str, scan_result: dict) -> dict:
    idx_a = text.find("B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG")
    idx_b = text.find("B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG")
    idx_end = text.find("Hai b\u00ean t\u1ef1 nguy\u1ec7n")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="structured_docx")


# ============================================================
# HAM CHINH
# ============================================================
def extract(filepath, *, scan_result: dict | None = None, preloaded_text: str | None = None):
    path = Path(filepath)
    use_ifilter_for_doc = path.suffix.lower() == ".doc"

    text = preloaded_text
    if text is None:
        text = read_docx(path, use_ifilter_for_doc=use_ifilter_for_doc)
        if use_ifilter_for_doc:
            text = _normalize_plain_text_for_extract(text)

    if scan_result is None:
        scan_result = _scan_contract_text(text, file_name=path.name)

    if use_ifilter_for_doc:
        return _build_payload_generic(path, text, scan_result, extract_mode="plain_text_doc")
    return _build_payload_generic(path, text, scan_result, extract_mode="structured_docx")

    text = read_docx(filepath)
    scan_result = scan_docx_for_contract_no(filepath)

    idx_a = text.find("BÊN CHUYỂN NHƯỢNG")
    idx_b = text.find("BÊN NHẬN CHUYỂN NHƯỢNG")
    idx_end = text.find("Hai bên tự nguyện")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}

    ten_hd = find_ten_hop_dong(text)
    tai_san = find_tai_san(text)
    so_cong_chung = find_so_cong_chung(text) or scan_result.get("contract_no", "")

    return {
        "web_form": {
            "ten_hop_dong": ten_hd,
            "ngay_cong_chung": find_ngay_cong_chung(text),
            "so_cong_chung": so_cong_chung,
            "nhom_hop_dong": guess_nhom_hd(ten_hd),
            "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
            "cong_chung_vien": find_ccv(text),
            "thu_ky": DEFAULT_THU_KY,
            "nguoi_yeu_cau": fmt_nguoi_yeu_cau(ben_b),
            "duong_su": fmt_duong_su(ben_a, ben_b),
            "tai_san": tai_san,
        },
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": os.path.abspath(filepath),
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
        },
    }


def main():
    if len(sys.argv) < 2:
        print("Cách dùng:")
        print('  python extract_contract.py "hop_dong.docx"')
        print('  python extract_contract.py "hop_dong.doc"')
        sys.exit(1)

    fp = sys.argv[1]
    if not os.path.exists(fp):
        print(f"Không tìm thấy file: {fp}")
        sys.exit(1)

    result = extract(fp)
    output = json.dumps(result, ensure_ascii=False, indent=2)
    print(output)

    json_path = fp.rsplit(".", 1)[0] + "_extracted.json"
    with open(json_path, "w", encoding="utf-8") as file_obj:
        file_obj.write(output)
    print(f"\n→ Đã lưu: {json_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
