"""
TRICH XUAT HOP DONG CONG CHUNG -> JSON CHO WEB FORM
===================================================
Cach dung:
    python extract_contract.py "hop_dong.docx"
    python extract_contract.py "hop_dong.doc"

Ket qua:
    - In JSON ra man hinh
    - Luu file .json cung thu muc

Cai dat:
    pip install python-docx
"""

from __future__ import annotations

import json
import os
import re
import sys
import unicodedata
from pathlib import Path
from typing import Iterable


# ============================================================
# CAU HINH MAC DINH — sua tai day cho phu hop van phong
# ============================================================
DEFAULT_CCV = "Phạm Minh Chi"
DEFAULT_THU_KY = "Nguyễn Nhật Minh"
CONTRACT_YEAR = "2026"
CONTRACT_NO_REGEX = re.compile(rf"\b(\d+/{re.escape(CONTRACT_YEAR)}/CCGD)\b", re.IGNORECASE)
CONTRACT_KEYWORDS = ("hợp đồng", "hop dong", "hđ", "hd")

NHOM_HD_MAP = {
    "chuyển nhượng": "Chuyển nhượng - Mua bán",
    "mua bán": "Chuyển nhượng - Mua bán",
    "tặng cho": "Tặng, cho tài sản",
    "tặng, cho": "Tặng, cho tài sản",
    "thế chấp": "Cầm cố - Thế chấp - Vay",
    "cầm cố": "Cầm cố - Thế chấp - Vay",
    "vay": "Cầm cố - Thế chấp - Vay",
    "ủy quyền": "Ủy quyền",
    "ủy quền": "Ủy quyền",
    "di chúc": "Di chúc",
    "đặt cọc": "Đặt cọc",
    "thỏa thuận": "Thoả thuận - Cam kết",
    "cam kết": "Thoả thuận - Cam kết",
    "góp vốn": "Góp vốn - Hợp tác",
    "hợp tác": "Góp vốn - Hợp tác",
    "thừa kế": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "khai nhận": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "phân chia": "Thừa kế (khai nhận - phân chia di sản thừa kế )",
    "từ chối nhận di sản": "Từ chối nhận di sản thừa kế",
    "phụ lục": "Phụ lục hợp đồng - văn bản sửa đổi",
    "thuê": "Thuê - Mượn tài sản",
    "mượn": "Thuê - Mượn tài sản",
}


CRITICAL_WEB_FORM_FIELDS = (
    "ten_hop_dong",
    "so_cong_chung",
    "nhom_hop_dong",
    "loai_tai_san",
    "tai_san",
)

DOC_KIND_TRANSFER = "transfer_contract"
DOC_KIND_ASSET_COMMITMENT = "asset_commitment"
DOC_KIND_TRANSFER_CANCELLATION = "transfer_cancellation"
DOC_KIND_GENERIC = "generic"


def _append_text_lines(lines: list[str], values: Iterable[str]) -> None:
    for raw in values:
        text = str(raw or "").strip()
        if text:
            lines.append(text)


def _append_table_lines(lines: list[str], tables) -> None:
    for table in tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text and p.text.strip()
                ).strip()
                if cell_text:
                    row_values.append(cell_text)
            if row_values:
                lines.append(" | ".join(row_values))


# ============================================================
# DOCX / DOC HELPERS
# ============================================================
def read_doc_via_ifilter(filepath) -> str:
    """Trich xuat text tu .doc bang Windows IFilter (cung co che voi Agent Ransack).
    WHY: Doc .doc tren Windows ma khong can Word/Office.
         Dung cho scan (tim so hop dong) — chi can plain text, khong can table structure.
    RISK: IFilter tra plain text thuan — mat cau truc bang va header/footer.
          Flow extract() cho .doc se chay theo plain text nay.
    REQUIRE: query.dll (co san tren Windows 7+, khong can cai them).
    """
    import ctypes

    abs_path = str(Path(filepath).resolve())
    ole32 = ctypes.windll.ole32
    ole32.CoInitialize(None)
    try:
        query = ctypes.windll.query
        # c_long (signed) de compare HRESULT am — WINFUNCTYPE(HRESULT) tu raise exception
        query.LoadIFilter.restype = ctypes.c_long
        query.LoadIFilter.argtypes = [
            ctypes.c_wchar_p,
            ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_void_p),
        ]
        p_filter = ctypes.c_void_p()
        hr = query.LoadIFilter(abs_path, None, ctypes.byref(p_filter))
        if hr != 0 or not p_filter.value:
            raise OSError(f"LoadIFilter that bai (hr=0x{hr & 0xFFFFFFFF:08X})")

        # IFilter VTBL layout (IUnknown 0-2, IFilter 3-7):
        #   3=Init, 4=GetChunk, 5=GetText, 6=GetValue, 7=BindRegion
        VTBL_INIT = 3
        VTBL_GET_CHUNK = 4
        VTBL_GET_TEXT = 5
        VTBL_RELEASE = 2

        vtbl_ptr = ctypes.cast(p_filter, ctypes.POINTER(ctypes.c_void_p))
        vtbl = ctypes.cast(vtbl_ptr[0], ctypes.POINTER(ctypes.c_void_p))

        # Init(grfFlags, cAttributes, aAttributes, pFlags) — c_long tranh auto-raise
        InitFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p,
            ctypes.c_uint, ctypes.c_ulong, ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_uint),
        )
        out_flags = ctypes.c_uint(0)
        InitFunc(vtbl[VTBL_INIT])(p_filter, 0, 0, None, ctypes.byref(out_flags))

        # STAT_CHUNK tren x64 Windows:
        #   offset  0: idChunk   (ULONG 4B)
        #   offset  4: breakType (enum  4B)
        #   offset  8: flags     (CHUNKSTATE 4B) — CHUNK_TEXT=0x1, CHUNK_VALUE=0x2
        #   offset 12: locale    (LCID  4B)
        #   offset 16: attribute (FULLPROPSPEC = GUID(16B) + PROPSPEC(4+4pad+8B) = 32B)
        #   offset 48: idChunkSource (4B), cwcStartSource (4B), cwcLenSource (4B)
        #   total: 60B — dung 64B buffer cho an toan tren moi alignment
        CHUNK_BUF_SIZE = 64
        CHUNK_FLAGS_OFFSET = 8
        CHUNK_TEXT = 0x1  # CHUNKSTATE: CHUNK_TEXT=1 (khong phai 2)

        GetChunkFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p, ctypes.c_char_p,
        )
        get_chunk_fn = GetChunkFunc(vtbl[VTBL_GET_CHUNK])

        BUF_SIZE = 4096
        GetTextFunc = ctypes.WINFUNCTYPE(
            ctypes.c_long, ctypes.c_void_p,
            ctypes.POINTER(ctypes.c_ulong), ctypes.c_wchar_p,
        )
        get_text_fn = GetTextFunc(vtbl[VTBL_GET_TEXT])

        ReleaseFunc = ctypes.WINFUNCTYPE(ctypes.c_ulong, ctypes.c_void_p)
        release_fn = ReleaseFunc(vtbl[VTBL_RELEASE])

        # HRESULT codes (signed 32-bit)
        FILTER_E_END_OF_CHUNKS = ctypes.c_long(0x80041700).value  # negative
        FILTER_E_NO_MORE_TEXT  = ctypes.c_long(0x80041701).value  # negative
        # FILTER_S_LAST_TEXT = 0x00041709 — SUCCESS: text duoc ghi VA day la phan cuoi cung
        # Phai capture text TRUOC khi check status nay, roi break
        FILTER_S_LAST_TEXT = 0x00041709  # positive (success with info)

        parts: list[str] = []
        chunk_buf = ctypes.create_string_buffer(CHUNK_BUF_SIZE)

        while True:
            hr_chunk = get_chunk_fn(p_filter, chunk_buf)
            if hr_chunk == FILTER_E_END_OF_CHUNKS:
                break
            if hr_chunk != 0:
                break

            flags = ctypes.c_uint.from_buffer_copy(chunk_buf, CHUNK_FLAGS_OFFSET).value
            if not (flags & CHUNK_TEXT):
                continue

            while True:
                buf_len = ctypes.c_ulong(BUF_SIZE)
                text_buf = ctypes.create_unicode_buffer(BUF_SIZE)
                hr_text = get_text_fn(p_filter, ctypes.byref(buf_len), text_buf)
                # Capture text TRUOC khi check status — FILTER_S_LAST_TEXT tra text hop le
                if buf_len.value > 0:
                    parts.append(text_buf.value[: buf_len.value])
                if hr_text == FILTER_E_NO_MORE_TEXT:
                    break
                if hr_text == FILTER_S_LAST_TEXT:
                    break  # text da duoc lay, day la phan cuoi cua chunk nay
                if hr_text != 0:
                    break

        release_fn(p_filter)
        return "".join(parts)
    finally:
        ole32.CoUninitialize()


def read_docx(filepath, *, use_ifilter_for_doc: bool = False):
    """Doc .docx hoac .doc -> text thuan.
    - .docx: dung python-docx truc tiep (nhanh).
    - .doc: dung IFilter/plain text tren Windows, khong can Word.
    """
    path = Path(filepath)
    if path.suffix.lower() == ".doc":
        return read_doc_via_ifilter(filepath)

    from docx import Document
    doc = Document(str(filepath))
    lines: list[str] = []
    _append_text_lines(lines, (p.text for p in doc.paragraphs))
    _append_table_lines(lines, doc.tables)
    for section in doc.sections:
        _append_text_lines(lines, (p.text for p in section.header.paragraphs))
        _append_text_lines(lines, (p.text for p in section.footer.paragraphs))
    return "\n".join(lines)


def _normalize_scan_contract_no(contract_no: str) -> str:
    return str(contract_no or "").strip().upper()


def _normalize_web_contract_no(contract_no: str) -> str:
    normalized = _normalize_scan_contract_no(contract_no)
    if normalized.endswith("/CCGD"):
        normalized = normalized[:-5]
    return normalized


def find_contract_numbers(text: str) -> list[str]:
    return [_normalize_scan_contract_no(match.group(1)) for match in CONTRACT_NO_REGEX.finditer(text or "")]


def has_contract_keyword(text: str, file_name: str = "") -> bool:
    haystack = f"{file_name}\n{text}".lower()
    return any(keyword in haystack for keyword in CONTRACT_KEYWORDS)


def _normalize_plain_text_for_extract(text: str) -> str:
    normalized = str(text or "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\x0b", "\n").replace("\t", " ").replace("\xa0", " ")
    normalized = unicodedata.normalize("NFC", normalized)
    lines = []
    for raw in normalized.split("\n"):
        line = re.sub(r"\s+", " ", raw).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _scan_contract_text(text: str, *, file_name: str = "") -> dict:
    matches = find_contract_numbers(text)
    if matches:
        contract_no = matches[-1]
        return {
            "is_contract": True,
            "contract_no": contract_no,
            "reason": f"Tim thay so CC: {contract_no}",
        }

    if has_contract_keyword(text, file_name):
        return {
            "is_contract": False,
            "contract_no": "",
            "reason": "Co keyword hop dong nhung khong thay so cong chung",
        }

    return {
        "is_contract": False,
        "contract_no": "",
        "reason": "Khong thay so cong chung",
    }


def scan_docx_for_contract_no(filepath, *, include_text: bool = False):
    """Tim so cong chung trong noi dung file .docx hoac .doc.
    .doc dung IFilter (nhanh) — chi can plain text de tim regex so cong chung.
    """
    path = Path(filepath)
    use_ifilter_for_doc = path.suffix.lower() == ".doc"
    try:
        text = read_docx(path, use_ifilter_for_doc=use_ifilter_for_doc)
        if use_ifilter_for_doc:
            text = _normalize_plain_text_for_extract(text)
    except Exception as exc:
        result = {
            "is_contract": False,
            "contract_no": "",
            "reason": f"Khong doc duoc file: {exc}",
        }
        if include_text:
            result["text"] = ""
        return result

    result = _scan_contract_text(text, file_name=path.name)
    if include_text:
        result["text"] = text
    return result


# ============================================================
# TRICH XUAT TUNG TRUONG
# ============================================================
def find_ten_hop_dong(text):
    return _detect_document_kind_and_title(text)[1]


def find_so_cong_chung(text):
    m = re.search(r"Số công chứng\s*(\d+)\s*/\s*(\d{4})\s*/\s*CCGD\b", text, re.IGNORECASE)
    if m:
        return _normalize_web_contract_no(f"{m.group(1)}/{m.group(2)}/CCGD")

    matches = find_contract_numbers(text)
    if matches:
        return _normalize_web_contract_no(matches[-1])
    return ""


def find_ngay_cong_chung(text):
    idx = text.find("LỜI CHỨNG")
    sub = text[idx:] if idx >= 0 else text
    m = re.search(r"ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})", sub)
    if m:
        return f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/{m.group(3)}"
    return ""


def find_ccv(text):
    m = re.search(r"Tôi,\s+([^,]+),\s+Công chứng viên", text)
    return m.group(1).strip() if m else DEFAULT_CCV


def find_persons(section):
    """Tim tat ca nguoi trong 1 doan (Ben A hoac Ben B)."""
    persons = []
    names = list(re.finditer(r"(Ông|[Bb]à)\s+([^;]+?)\s*;", section))

    for idx, nm in enumerate(names):
        start = nm.start()
        end = names[idx + 1].start() if idx + 1 < len(names) else len(section)
        chunk = section[start:end]

        person = {
            "gioi_tinh": nm.group(1).capitalize(),
            "ho_ten": nm.group(2).strip(),
        }

        m = re.search(r"Sinh ngày:?\s*(\d{2}/\d{2}/\d{4})", chunk)
        person["ngay_sinh"] = m.group(1) if m else ""

        m = re.search(r"(?:Căn cước|CCCD|CMND)(?:\s+công dân)?\s*(?:số:?\s*)(\d+)", chunk)
        person["cccd"] = m.group(1) if m else ""

        m = re.search(r"do\s+(.+?)\s+cấp ngày", chunk, re.DOTALL)
        person["noi_cap"] = re.sub(r"\s+", " ", m.group(1)).strip() if m else ""

        m = re.search(r"cấp ngày\s*(\d{2}/\d{2}/\d{4})", chunk)
        person["ngay_cap_cccd"] = m.group(1) if m else ""

        persons.append(person)

    return persons


def find_dia_chi(section):
    m = re.search(r"cùng (?:cư trú|thường trú) tại:?\s*(.+?)\.", section, re.DOTALL)
    return re.sub(r"\s+", " ", m.group(1)).strip() if m else ""


def find_tai_san(text):
    doc_kind, _ = _detect_document_kind_and_title(text)
    return _find_tai_san_by_kind(text, doc_kind)


def _fold_text(text: str) -> str:
    normalized = unicodedata.normalize("NFD", str(text or ""))
    normalized = normalized.replace("\u0111", "d").replace("\u0110", "D")
    return "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn").lower()


def _fold_text_with_index_map(text: str) -> tuple[str, list[int]]:
    source = str(text or "")
    folded_chars: list[str] = []
    index_map: list[int] = []
    for idx, char in enumerate(source):
        normalized = unicodedata.normalize("NFD", char)
        normalized = normalized.replace("\u0111", "d").replace("\u0110", "D")
        for part in normalized:
            if unicodedata.category(part) == "Mn":
                continue
            folded_chars.append(part.lower())
            index_map.append(idx)
    return "".join(folded_chars), index_map


def _sentence_case_vn(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip(" .:-;\n\t")
    if not cleaned:
        return ""
    return cleaned[:1].upper() + cleaned[1:].lower()


def _clean_title_line(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip(" .:-;\n\t")
    if not cleaned:
        return ""
    cleaned = re.sub(
        r"\s+(?:này\s+)?được\s+giao\s+kết\s+bởi\s*:?\s*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(
        r"\s+(?:nay\s+)?duoc\s+giao\s+ket\s+boi\s*:?\s*$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    return cleaned.strip(" .:-;\n\t")


def _find_title_line(text: str) -> str:
    for raw in str(text or "").split("\n")[:20]:
        line = _clean_title_line(raw)
        if not line:
            continue
        folded = _fold_text(line)
        if (
            folded.startswith("hop dong ")
            or folded.startswith("van ban ")
            or folded.startswith("thoa thuan ")
            or "cam ket" in folded
        ):
            return line
    return ""


def _detect_document_kind_and_title(text: str, *, file_name: str = "") -> tuple[str, str]:
    title_line = _clean_title_line(_find_title_line(text))
    title_fold = _fold_text(title_line)
    combined_fold = _fold_text(f"{file_name}\n{text}")
    header_fold = _fold_text("\n".join(str(text or "").split("\n")[:40]))
    title_or_header_fold = f"{title_fold}\n{header_fold}"

    if "cam ket tai san rieng" in title_fold or "cam ket tai san rieng" in combined_fold:
        return DOC_KIND_ASSET_COMMITMENT, "Văn bản cam kết tài sản riêng"

    is_transfer_cancellation = (
        ("huy bo" in title_fold or "huy hop dong" in title_fold)
        and (
            "hop dong chuyen nhuong" in title_or_header_fold
            or "chuyen nhuong quyen su dung dat" in title_or_header_fold
        )
    )
    if is_transfer_cancellation:
        if any(
            marker in combined_fold
            for marker in (
                "tai san gan lien voi dat",
                "nha o va quyen su dung dat o",
                "mua ban nha o",
            )
        ):
            return (
                DOC_KIND_TRANSFER_CANCELLATION,
                "Văn bản thỏa thuận về việc hủy bỏ hợp đồng chuyển nhượng quyền sử dụng đất và tài sản gắn liền với đất",
            )
        return (
            DOC_KIND_TRANSFER_CANCELLATION,
            "Văn bản thỏa thuận về việc hủy bỏ hợp đồng chuyển nhượng quyền sử dụng đất",
        )

    if "hop dong chuyen nhuong" in title_or_header_fold or "hop dong chuyen nhuong" in combined_fold:
        if any(
            marker in combined_fold
            for marker in (
                "tai san gan lien voi dat",
                "nha o va quyen su dung dat o",
                "va nha o",
                "va tai san",
            )
        ):
            return DOC_KIND_TRANSFER, "Hợp đồng chuyển nhượng quyền sử dụng đất"
        return DOC_KIND_TRANSFER, "Hợp đồng chuyển nhượng quyền sử dụng đất"

    if title_line:
        return DOC_KIND_GENERIC, _sentence_case_vn(title_line)
    return DOC_KIND_GENERIC, ""


def _extract_block_by_patterns(text: str, start_patterns: Iterable[str], end_patterns: Iterable[str]) -> str:
    match = _find_first_pattern_match_generic(text, start_patterns)
    if not match:
        return ""

    start_idx = int(match[0])
    end_idx = len(text)
    for pattern in end_patterns:
        end_match = re.search(pattern, text[start_idx:], re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not end_match:
            continue
        candidate_end = start_idx + end_match.start()
        if candidate_end > start_idx and candidate_end < end_idx:
            end_idx = candidate_end

    raw = text[start_idx:end_idx].strip()
    if not raw:
        return ""
    raw = unicodedata.normalize("NFC", raw)
    lines = [re.sub(r"\s+", " ", line).strip(" ;,") for line in raw.split("\n") if line.strip()]
    return "\n".join(lines).strip()


def _extract_block_by_folded_patterns(text: str, start_patterns: Iterable[str], end_patterns: Iterable[str]) -> str:
    source = str(text or "")
    if not source:
        return ""

    folded_text, index_map = _fold_text_with_index_map(source)
    if not folded_text or not index_map:
        return ""

    best_match = None
    for pattern in start_patterns:
        match = re.search(pattern, folded_text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not match:
            continue
        if best_match is None or match.start() < best_match.start():
            best_match = match

    if not best_match:
        return ""

    start_fold_idx = best_match.start()
    start_idx = index_map[start_fold_idx]
    end_idx = len(source)

    for pattern in end_patterns:
        end_match = re.search(pattern, folded_text[best_match.end():], re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not end_match:
            continue
        candidate_fold_idx = best_match.end() + end_match.start()
        if candidate_fold_idx <= start_fold_idx:
            continue
        candidate_end = index_map[candidate_fold_idx]
        if candidate_end > start_idx and candidate_end < end_idx:
            end_idx = candidate_end

    raw = source[start_idx:end_idx].strip()
    if not raw:
        return ""
    raw = unicodedata.normalize("NFC", raw)
    lines = [re.sub(r"\s+", " ", line).strip(" ;,") for line in raw.split("\n") if line.strip()]
    return "\n".join(lines).strip()


def _find_tai_san_generic(text: str) -> str:
    return _extract_block_by_patterns(
        text,
        (
            r"(?m)^\s*(?:1\.1\s*)?\u0110\u1ed1i t\u01b0\u1ee3ng c\u1ee7a H\u1ee3p \u0111\u1ed3ng n\u00e0y l\u00e0\b",
            r"(?m)^\s*\u0110I\u1ec0U\s*1\b",
        ),
        (
            r"(?m)^\s*1\.2\b",
            r"(?m)^\s*B\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
            r"(?m)^\s*\u0110I\u1ec0U\s*2\b",
            r"(?m)^\s*Gi\u00e1 chuy\u1ec3n nh\u01b0\u1ee3ng\b",
            r"(?m)^\s*Ph\u01b0\u01a1ng th\u1ee9c thanh to\u00e1n\b",
        ),
    )


def _find_tai_san_qsdd_common(text: str) -> str:
    return _extract_block_by_folded_patterns(
        text,
        (
            r"quyen su dung dat(?: va tai san gan lien voi dat)?[^\n]{0,250}?co dia chi tai\s*:?",
            r"nha o va quyen su dung dat o[^\n]{0,250}?co dia chi tai\s*:?",
            r"chu so huu nha o va quyen su dung dat o[^\n]{0,250}?co dia chi tai\s*:?",
        ),
        (
            r"(?m)^\s*1\.2\b",
            r"(?m)^\s*dieu\s*2\b",
            r"(?m)^\s*bang hop dong nay\b",
            r"(?m)^\s*bang van ban nay chung toi xac dinh\b",
            r"(?m)^\s*hai vo chong chung toi cam ket\b",
            r"(?m)^\s*hai vo chong chung toi cam doan\b",
            r"(?m)^\s*chung toi cong nhan\b",
            r"(?m)^\s*gia chuyen nhuong\b",
            r"(?m)^\s*phuong thuc thanh toan\b",
            r"\bva duoc cong chung vien\b",
            r"\bva duoc[\s\S]{0,120}?chung nhan\b",
            r"\bso cong chung\b",
        ),
    )


def _find_tai_san_by_kind(text: str, doc_kind: str) -> str:
    common_qsdd_block = _find_tai_san_qsdd_common(text)
    if common_qsdd_block:
        return common_qsdd_block

    if doc_kind == DOC_KIND_TRANSFER:
        return _extract_block_by_patterns(
            text,
            (r"(?m)^\s*(?:1\.1\s*)?\u0110\u1ed1i t\u01b0\u1ee3ng c\u1ee7a H\u1ee3p \u0111\u1ed3ng n\u00e0y l\u00e0\b",),
            (
                r"(?m)^\s*1\.2\b",
                r"(?m)^\s*B\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
                r"(?m)^\s*\u0110I\u1ec0U\s*2\b",
                r"(?m)^\s*Gi\u00e1 chuy\u1ec3n nh\u01b0\u1ee3ng\b",
                r"(?m)^\s*Ph\u01b0\u01a1ng th\u1ee9c thanh to\u00e1n\b",
            ),
        )

    if doc_kind == DOC_KIND_ASSET_COMMITMENT:
        return _extract_block_by_patterns(
            text,
            (
                r"(?m)^\s*(?:\u00d4ng|B\u00e0)[^\n]{0,200}?hi\u1ec7n \u0111ang s\u1edf h\u1eefu\s+T\u00e0i\s*S\u1ea3n\s+l\u00e0\b",
                r"(?m)^\s*T\u00e0i\s*S\u1ea3n\s+l\u00e0\b",
            ),
            (
                r"(?m)^\s*B\u1eb1ng v\u0103n b\u1ea3n n\u00e0y ch\u00fang t\u00f4i x\u00e1c \u0111\u1ecbnh\b",
                r"(?m)^\s*Hai v\u1ee3 ch\u1ed3ng ch\u00fang t\u00f4i cam \u0111oan\b",
                r"(?m)^\s*Ch\u00fang t\u00f4i c\u00f4ng nh\u1eadn\b",
            ),
        )

    if doc_kind == DOC_KIND_TRANSFER_CANCELLATION:
        return _extract_block_by_patterns(
            text,
            (
                r"(?:mua b\u00e1n nh\u00e0 \u1edf v\u00e0\s+)?chuy\u1ec3n nh\u01b0\u1ee3ng quy\u1ec1n s\u1eed d\u1ee5ng \u0111\u1ea5t(?:\s+v\u00e0\s+t\u00e0i\s+s\u1ea3n\s+g\u1eafn\s+li\u1ec1n\s+v\u1edbi\s+\u0111\u1ea5t)?[^.\n]{0,200}?c\u00f3 \u0111\u1ecba ch\u1ec9 t\u1ea1i:\s*",
            ),
            (
                r"\bv\u00e0 \u0111\u01b0\u1ee3c C\u00f4ng ch\u1ee9ng vi\u00ean\b",
                r"\bv\u00e0 \u0111\u01b0\u1ee3c[^.\n]{0,120}?ch\u1ee9ng nh\u1eadn\b",
                r"\bs\u1ed1 c\u00f4ng ch\u1ee9ng\b",
            ),
        )

    return _find_tai_san_generic(text)


def guess_nhom_hd(ten_hd):
    lower = ten_hd.lower()
    for kw, nhom in NHOM_HD_MAP.items():
        if kw in lower:
            return nhom
    return "Khác"


LAND_CERTIFICATE_HEADING_NOISE_PATTERNS = (
    r"giay chung nhan quyen su dung dat\s*,?\s*quyen so huu nha o va tai san khac gan lien voi dat",
    r"giay chung nhan quyen su dung dat\s*,?\s*quyen so huu tai san gan lien voi dat",
    r"giay chung nhan quyen su dung dat\s*,?\s*quyen so huu nha o",
)

LAND_WITH_ASSET_TITLE_MARKERS = (
    "quyen su dung dat va tai san gan lien voi dat",
    "nha o va quyen su dung dat o",
    "mua ban nha o",
)

LAND_ATTACHED_ASSET_DETAIL_PATTERNS = (
    r"\bnha o\b\s*:",
    r"\bquyen so huu nha o\b",
    r"\bnha o va quyen su dung dat o\b",
    r"\bcong trinh xay dung\b\s*:",
    r"\btai san gan lien voi dat\b\s*:",
    r"\btai san khac gan lien voi dat\b\s*:",
    r"\bdien tich xay dung\b",
    r"\bdien tich san\b",
    r"\bket cau\b",
    r"\bso tang\b",
    r"\bhinh thuc so huu\b",
    r"\bcap \(hang\)\b",
    r"\bcap cong trinh\b",
)


def _strip_land_certificate_heading_noise(text_folded: str) -> str:
    cleaned = str(text_folded or "")
    for pattern in LAND_CERTIFICATE_HEADING_NOISE_PATTERNS:
        cleaned = re.sub(pattern, " ", cleaned, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", cleaned).strip()


def _has_attached_land_asset_details(text_folded: str) -> bool:
    cleaned = _strip_land_certificate_heading_noise(text_folded)
    return any(
        re.search(pattern, cleaned, re.IGNORECASE)
        for pattern in LAND_ATTACHED_ASSET_DETAIL_PATTERNS
    )


def guess_loai_tai_san(tai_san, ten_hd):
    tai_san_fold = _fold_text(tai_san)
    title_fold = _fold_text(ten_hd)
    combined = f"{tai_san_fold} {title_fold}".strip()

    if "quyen su dung dat" in combined:
        if any(marker in title_fold for marker in LAND_WITH_ASSET_TITLE_MARKERS):
            return "Đất đai có tài sản"
        if _has_attached_land_asset_details(tai_san_fold):
            return "Đất đai có tài sản"
        return "Đất đai không có tài sản"

    keywords = [
        ("o to", "Ô tô"),
        ("xe may", "Xe máy"),
        ("tau", "Tàu, thuyền"),
        ("thuyen", "Tàu, thuyền"),
        ("so tiet kiem", "Sổ tiết kiệm"),
        ("co phan", "Cổ phần"),
        ("co phieu", "Cổ phiếu"),
        ("chung khoan", "Chứng khoán"),
        ("the atm", "Thẻ ATM"),
    ]
    for kw, loai in keywords:
        if kw in combined:
            return loai
    return "Tài sản khác"


# ============================================================
# FORMAT CHO 3 O CKEDITOR TREN WEB
# ============================================================
def _normalize_person_block_lines(person: dict) -> list[str]:
    raw_text = unicodedata.normalize("NFC", str(person.get("raw_text") or "")).strip()
    lines = [re.sub(r"\s+", " ", line).strip(" ;") for line in raw_text.split("\n") if line.strip()]
    if not lines:
        first_line = " ".join(
            part
            for part in (
                f"{person.get('gioi_tinh', '')} {person.get('ho_ten', '')}".strip(),
                f"Sinh ngày: {person['ngay_sinh']}" if person.get("ngay_sinh") else "",
            )
            if part
        ).strip()
        if first_line:
            lines.append(first_line)
        if person.get("cccd"):
            id_line = f"Căn cước số: {person['cccd']}"
            if person.get("noi_cap") and person.get("ngay_cap_cccd"):
                id_line += f" do {person['noi_cap']} cấp ngày {person['ngay_cap_cccd']}"
            lines.append(id_line)

    if lines:
        lines[0] = re.sub(
            r"^(?:[-*]\s*)?(?:(?:Người|người)\s+(?:vợ|chồng)\s*[–-]\s*)?",
            "",
            lines[0],
        ).strip()
        lines[0] = re.sub(r"^(?:\d+\.\s*)+", "", lines[0]).strip()
        lines[0] = re.sub(r"^(?:Và|và)\s+(?:vợ|chồng)\s+là\s+", "", lines[0]).strip()
        if lines[0]:
            lines[0] = lines[0][:1].upper() + lines[0][1:]

    address_line = str(person.get("dia_chi_line") or "").strip(" ;")
    if not address_line and person.get("dia_chi"):
        address_line = f"Thường trú tại: {person['dia_chi']}"
    if address_line and not any("tru tai" in _fold_text(line) for line in lines):
        lines.append(address_line)
    return lines


def _format_person_block(person: dict, *, index: int | None = None) -> str:
    lines = _normalize_person_block_lines(person)
    if not lines:
        return ""
    if index is not None:
        lines[0] = f"{index}. {lines[0]}"
    return "\n".join(lines).strip()


def fmt_nguoi_yeu_cau(ben_b):
    """Nguoi yeu cau CC = nguoi dau tien ben B."""
    if not ben_b.get("nguoi"):
        return ""
    return _format_person_block(ben_b["nguoi"][0])


def fmt_duong_su(ben_a, ben_b):
    """Duong su = toan bo ben A + ben B."""
    lines = ["BÊN A (Bên chuyển nhượng):"]
    for idx, person in enumerate(ben_a.get("nguoi", []), 1):
        block = _format_person_block(person, index=idx)
        if block:
            lines.append(block)
            lines.append("")

    lines.append("BÊN B (Bên nhận chuyển nhượng):")
    for idx, person in enumerate(ben_b.get("nguoi", []), 1):
        block = _format_person_block(person, index=idx)
        if block:
            lines.append(block)
            lines.append("")

    return "\n".join(lines).strip()


def get_missing_web_form_fields(web_form: dict, *, file_hop_dong: str = "") -> list[str]:
    values = dict(web_form or {})
    values["file_hop_dong"] = file_hop_dong
    return [field for field in CRITICAL_WEB_FORM_FIELDS if not str(values.get(field) or "").strip()]


def _find_first_pattern_index(text: str, patterns: Iterable[str], *, start: int = 0) -> int:
    haystack = text[start:]
    best_idx = -1
    for pattern in patterns:
        match = re.search(pattern, haystack, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not match:
            continue
        idx = start + match.start()
        if best_idx < 0 or idx < best_idx:
            best_idx = idx
    return best_idx


def _extract_plain_text_persons(section: str) -> list[dict]:
    return [dict(entry) for entry in _extract_plain_text_person_entries(section)]


def _extract_plain_text_person_entries(section: str) -> list[dict]:
    section = unicodedata.normalize("NFC", str(section or ""))
    header_re = re.compile(
        r"(?i)^\s*(?:[-*]\s*)?(?:(?:ng\u01b0\u1eddi\s+)?(?:v\u1ee3|ch\u1ed3ng)\s*[–-]\s*)?"
        r"(?:\d+\.\s*)?(?:v\u00e0\s+(?:v\u1ee3|ch\u1ed3ng)\s+l\u00e0\s+)?"
        r"(?:ng\u01b0\u1eddi\s+(?:v\u1ee3|ch\u1ed3ng)\s+)?(?P<title>\u00f4ng|b\u00e0)\s*:?\s*"
        r"(?P<name>.+?)(?=\s+Sinh ng\u00e0y:|[;]|$)"
    )
    birth_re = re.compile(r"(?i)Sinh ng\u00e0y:?\s*(\d{1,2}/\d{1,2}/\d{4})")
    id_re = re.compile(r"(?i)(?:C\u0103n c\u01b0\u1edbc(?:\s+c\u00f4ng d\u00e2n)?|CCCD|CMND)\s*(?:s\u1ed1)?\s*:?\s*(\d+)")
    issue_place_re = re.compile(r"(?i)\bdo\s+(.+?)\s+c\u1ea5p ng\u00e0y", re.DOTALL)
    issue_date_re = re.compile(r"(?i)c\u1ea5p ng\u00e0y\s*(\d{1,2}/\d{1,2}/\d{4})")
    address_re = re.compile(
        r"(?im)^(?:C\u1ea3 hai \u00f4ng b\u00e0\s+)?(?:C\u00f9ng\s+)?(?:N\u01a1i\s+)?"
        r"(?:th\u01b0\u1eddng tr\u00fa|c\u01b0 tr\u00fa)\s+t\u1ea1i:\s*(.+)$"
    )

    lines = [re.sub(r"\s+", " ", line).strip() for line in section.split("\n") if line.strip()]
    headers: list[tuple[int, re.Match[str]]] = []
    address_lines: list[tuple[int, str, str]] = []
    for idx, line in enumerate(lines):
        match = header_re.search(line)
        if match:
            headers.append((idx, match))
        address_match = address_re.search(line)
        if address_match:
            address_lines.append((idx, line.strip(" ;"), re.sub(r"\s+", " ", address_match.group(1)).strip(" .;")))

    persons = []
    for header_idx, (line_idx, match) in enumerate(headers):
        next_idx = headers[header_idx + 1][0] if header_idx + 1 < len(headers) else len(lines)
        chunk_lines = [line.strip(" ;") for line in lines[line_idx:next_idx] if line.strip()]
        chunk = "\n".join(chunk_lines)
        title = match.group("title").strip().lower()
        person = {
            "gioi_tinh": "\u00d4ng" if title.startswith("\u00f4") else "B\u00e0",
            "ho_ten": re.sub(r"\s+", " ", match.group("name")).strip(" ;,."),
            "_line_idx": line_idx,
            "raw_text": chunk,
        }
        birth_match = birth_re.search(chunk)
        person["ngay_sinh"] = birth_match.group(1) if birth_match else ""
        id_match = id_re.search(chunk)
        person["cccd"] = id_match.group(1) if id_match else ""
        issue_place_match = issue_place_re.search(chunk)
        person["noi_cap"] = re.sub(r"\s+", " ", issue_place_match.group(1)).strip(" ;,.") if issue_place_match else ""
        issue_date_match = issue_date_re.search(chunk)
        person["ngay_cap_cccd"] = issue_date_match.group(1) if issue_date_match else ""
        address_match = next((item for item in address_lines if line_idx <= item[0] < next_idx), None)
        person["dia_chi"] = address_match[2] if address_match else ""
        person["dia_chi_line"] = address_match[1] if address_match else ""
        name_fold = _fold_text(person["ho_ten"])
        if "hien dang so huu" in name_fold or "tai san la" in name_fold:
            continue
        if not birth_match and not id_match:
            continue
        persons.append(person)

    for person in persons:
        if not person.get("dia_chi"):
            fallback = next((item for item in address_lines if item[0] > person["_line_idx"]), None)
            if fallback:
                person["dia_chi"] = fallback[2]
                person["dia_chi_line"] = fallback[1]
                raw_text = str(person.get("raw_text") or "").strip()
                if fallback[1] not in raw_text:
                    person["raw_text"] = f"{raw_text}\n{fallback[1]}".strip()
        person.pop("_line_idx", None)
    return persons


def _extract_plain_text_address(section: str) -> str:
    section = unicodedata.normalize("NFC", str(section or ""))
    match = re.search(
        r"(?im)^(?:C\u1ea3 hai \u00f4ng b\u00e0\s+)?(?:C\u00f9ng\s+)?(?:N\u01a1i\s+)?"
        r"(?:th\u01b0\u1eddng tr\u00fa|c\u01b0 tr\u00fa)\s+t\u1ea1i:\s*(.+)$",
        section,
    )
    if not match:
        return ""
    return re.sub(r"\s+", " ", match.group(1)).strip(" .;")


def _extract_plain_text_parties(text: str) -> tuple[dict, dict]:
    ben_a_patterns = (r"\b(?:I\.\s*)?B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG\b", r"\bB\u00caN A\b")
    ben_b_patterns = (r"\b(?:II\.\s*)?B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG\b", r"\bB\u00caN B\b")
    end_patterns = (
        r"\bC\u00e1c b\u00ean \u0111\u00e3 t\u1ef1 nguy\u1ec7n\b",
        r"\bHai b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bB\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
        r"\b\u0110I\u1ec0U\s*1\b",
        r"\bL\u1edcI CH\u1ee8NG\b",
    )

    idx_a = _find_first_pattern_index(text, ben_a_patterns)
    idx_b = _find_first_pattern_index(text, ben_b_patterns, start=max(idx_a, 0))
    idx_end = _find_first_pattern_index(text, end_patterns, start=max(idx_b, 0)) if idx_b >= 0 else -1

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end > idx_b else text[idx_b:] if idx_b >= 0 else ""

    ben_a = {
        "nguoi": _extract_plain_text_persons(ben_a_section),
        "dia_chi": _extract_plain_text_address(ben_a_section),
    }
    ben_b = {
        "nguoi": _extract_plain_text_persons(ben_b_section),
        "dia_chi": _extract_plain_text_address(ben_b_section),
    }
    return ben_a, ben_b


def _find_first_pattern_match_generic(text: str, patterns: Iterable[str], *, start: int = 0):
    haystack = text[start:]
    best_match = None
    for pattern in patterns:
        match = re.search(pattern, haystack, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if not match:
            continue
        idx = start + match.start()
        end = start + match.end()
        if best_match is None or idx < best_match[0]:
            best_match = (idx, end)
    return best_match


def _extract_commitment_parties(text: str) -> tuple[dict, dict]:
    source = unicodedata.normalize("NFC", str(text or ""))
    lines = [re.sub(r"\s+", " ", line).strip() for line in source.split("\n") if line.strip()]
    if not lines:
        return {"nguoi": [], "dia_chi": ""}, {"nguoi": [], "dia_chi": ""}

    start_idx = 0
    for idx, line in enumerate(lines):
        if "chung toi gom" in _fold_text(line):
            start_idx = idx + 1
            break

    stop_prefixes = (
        "chung toi la vo chong",
        "nay chung toi lap van ban nay",
        "hien nay",
        "bang van ban nay",
        "loi chung",
    )

    preamble_lines: list[str] = []
    for line in lines[start_idx:]:
        folded = _fold_text(line)
        if any(folded.startswith(prefix) for prefix in stop_prefixes):
            break
        preamble_lines.append(line)

    entries = _extract_plain_text_person_entries("\n".join(preamble_lines))
    ben_a = {"nguoi": [], "dia_chi": ""}
    ben_b = {"nguoi": [], "dia_chi": ""}

    if entries:
        first = dict(entries[0])
        ben_a["dia_chi"] = first.pop("dia_chi", "")
        ben_a["nguoi"] = [first]
    if len(entries) > 1:
        second = dict(entries[1])
        ben_b["dia_chi"] = second.pop("dia_chi", "")
        ben_b["nguoi"] = [second]
    return ben_a, ben_b


def _extract_parties_generic(text: str) -> tuple[dict, dict]:
    doc_kind, _ = _detect_document_kind_and_title(text)
    if doc_kind == DOC_KIND_ASSET_COMMITMENT:
        ben_a, ben_b = _extract_commitment_parties(text)
        if ben_a.get("nguoi") or ben_b.get("nguoi"):
            return ben_a, ben_b

    ben_a_patterns = (
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHO VAY\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN CHO THU\u00ca\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN (?:UY|U\u1ef6|\u1ee6Y) QUY\u1ec0N\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN T\u1eb6NG CHO\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN B\u00c1N\b",
        r"(?m)^\s*(?:I\.\s*)?B\u00caN A\b",
    )
    ben_b_patterns = (
        r"(?m)^\s*(?:II\.\s*)?B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN VAY\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN THU\u00ca\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN \u0110\u01af\u1ee2C (?:UY|U\u1ef6|\u1ee6Y) QUY\u1ec0N\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN NH\u1eacN T\u1eb6NG CHO\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN MUA\b",
        r"(?m)^\s*(?:II\.\s*)?B\u00caN B\b",
    )
    preamble_patterns = (
        r"\bCh\u00fang t\u00f4i g\u1ed3m c\u00f3\s*:\s*",
        r"\bCh\u00fang t\u00f4i g\u1ed3m\s*:\s*",
    )
    end_patterns = (
        r"\bC\u00e1c b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bC\u00e1c b\u00ean \u0111\u00e3 t\u1ef1 nguy\u1ec7n\b",
        r"\bHai b\u00ean t\u1ef1 nguy\u1ec7n\b",
        r"\bB\u1eb1ng H\u1ee3p \u0111\u1ed3ng n\u00e0y\b",
        r"\b\u0110I\u1ec0U\s*1\b",
        r"\bL\u1edcI CH\u1ee8NG\b",
    )

    match_a = _find_first_pattern_match_generic(text, ben_a_patterns)
    match_b = _find_first_pattern_match_generic(text, ben_b_patterns)
    preamble_match = _find_first_pattern_match_generic(text, preamble_patterns)

    if match_a and match_b and match_a[0] > match_b[0]:
        match_a = None

    if match_a:
        idx_a_start = int(match_a[0])
    elif preamble_match and match_b:
        idx_a_start = int(preamble_match[1])
    else:
        idx_a_start = 0

    idx_b_start = int(match_b[0]) if match_b else -1
    idx_end = _find_first_pattern_index(text, end_patterns, start=max(idx_b_start, idx_a_start))

    if idx_b_start >= 0:
        ben_a_section = text[idx_a_start:idx_b_start]
        ben_b_section = text[idx_b_start:idx_end] if idx_end > idx_b_start else text[idx_b_start:]
    else:
        ben_a_section = text[idx_a_start:idx_end] if idx_end > idx_a_start else text[idx_a_start:]
        ben_b_section = ""

    ben_a = {
        "nguoi": _extract_plain_text_persons(ben_a_section),
        "dia_chi": _extract_plain_text_address(ben_a_section),
    }
    ben_b = {
        "nguoi": _extract_plain_text_persons(ben_b_section),
        "dia_chi": _extract_plain_text_address(ben_b_section),
    }
    return ben_a, ben_b


def _fmt_duong_su_generic(ben_a, ben_b):
    lines = ["BÊN A:"]
    for idx, person in enumerate(ben_a.get("nguoi", []), 1):
        block = _format_person_block(person, index=idx)
        if block:
            lines.append(block)
            lines.append("")

    lines.append("BÊN B:")
    for idx, person in enumerate(ben_b.get("nguoi", []), 1):
        block = _format_person_block(person, index=idx)
        if block:
            lines.append(block)
            lines.append("")
    return "\n".join(lines).strip()


def _build_payload_generic(filepath, text: str, scan_result: dict, *, extract_mode: str) -> dict:
    doc_kind, ten_hd = _detect_document_kind_and_title(text, file_name=Path(filepath).name)
    ben_a, ben_b = _extract_parties_generic(text)
    tai_san = _find_tai_san_by_kind(text, doc_kind)
    so_cong_chung = _normalize_web_contract_no(find_so_cong_chung(text) or scan_result.get("contract_no", ""))
    if doc_kind == DOC_KIND_ASSET_COMMITMENT and ben_a.get("nguoi"):
        nguoi_yeu_cau_party = ben_a
    else:
        nguoi_yeu_cau_party = ben_b if ben_b.get("nguoi") else ben_a
    web_form = {
        "ten_hop_dong": ten_hd,
        "ngay_cong_chung": find_ngay_cong_chung(text),
        "so_cong_chung": so_cong_chung,
        "nhom_hop_dong": guess_nhom_hd(ten_hd),
        "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
        "cong_chung_vien": find_ccv(text),
        "thu_ky": DEFAULT_THU_KY,
        "nguoi_yeu_cau": fmt_nguoi_yeu_cau(nguoi_yeu_cau_party),
        "duong_su": _fmt_duong_su_generic(ben_a, ben_b),
        "tai_san": tai_san,
    }
    file_goc = os.path.abspath(filepath)
    missing_fields = get_missing_web_form_fields(web_form, file_hop_dong=file_goc)
    return {
        "web_form": web_form,
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": file_goc,
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
            "document_kind": doc_kind,
            "extract_mode": extract_mode,
            "extract_is_partial": bool(missing_fields),
            "missing_web_form_fields": missing_fields,
        },
    }


def _build_payload(filepath, text: str, scan_result: dict, ben_a: dict, ben_b: dict, *, extract_mode: str) -> dict:
    doc_kind, ten_hd = _detect_document_kind_and_title(text, file_name=Path(filepath).name)
    tai_san = _find_tai_san_by_kind(text, doc_kind)
    so_cong_chung = _normalize_web_contract_no(find_so_cong_chung(text) or scan_result.get("contract_no", ""))
    if doc_kind == DOC_KIND_ASSET_COMMITMENT and ben_a.get("nguoi"):
        nguoi_yeu_cau_party = ben_a
    else:
        nguoi_yeu_cau_party = ben_b if ben_b.get("nguoi") else ben_a
    web_form = {
        "ten_hop_dong": ten_hd,
        "ngay_cong_chung": find_ngay_cong_chung(text),
        "so_cong_chung": so_cong_chung,
        "nhom_hop_dong": guess_nhom_hd(ten_hd),
        "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
        "cong_chung_vien": find_ccv(text),
        "thu_ky": DEFAULT_THU_KY,
        "nguoi_yeu_cau": fmt_nguoi_yeu_cau(nguoi_yeu_cau_party),
        "duong_su": fmt_duong_su(ben_a, ben_b),
        "tai_san": tai_san,
    }
    file_goc = os.path.abspath(filepath)
    missing_fields = get_missing_web_form_fields(web_form, file_hop_dong=file_goc)
    return {
        "web_form": web_form,
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": file_goc,
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
            "document_kind": doc_kind,
            "extract_mode": extract_mode,
            "extract_is_partial": bool(missing_fields),
            "missing_web_form_fields": missing_fields,
        },
    }


def _extract_structured_text_payload(filepath, text: str, scan_result: dict) -> dict:
    idx_a = text.find("BÃŠN CHUYá»‚N NHÆ¯á»¢NG")
    idx_b = text.find("BÃŠN NHáº¬N CHUYá»‚N NHÆ¯á»¢NG")
    idx_end = text.find("Hai bÃªn tá»± nguyá»‡n")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="structured_docx")


def _extract_plain_text_payload(filepath, text: str, scan_result: dict) -> dict:
    ben_a, ben_b = _extract_plain_text_parties(text)
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="plain_text_doc")


def _extract_structured_text_payload_v2(filepath, text: str, scan_result: dict) -> dict:
    idx_a = text.find("B\u00caN CHUY\u1ec2N NH\u01af\u1ee2NG")
    idx_b = text.find("B\u00caN NH\u1eacN CHUY\u1ec2N NH\u01af\u1ee2NG")
    idx_end = text.find("Hai b\u00ean t\u1ef1 nguy\u1ec7n")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}
    return _build_payload(filepath, text, scan_result, ben_a, ben_b, extract_mode="structured_docx")


# ============================================================
# HAM CHINH
# ============================================================
def extract(filepath, *, scan_result: dict | None = None, preloaded_text: str | None = None):
    path = Path(filepath)
    use_ifilter_for_doc = path.suffix.lower() == ".doc"

    text = preloaded_text
    if text is None:
        text = read_docx(path, use_ifilter_for_doc=use_ifilter_for_doc)
        if use_ifilter_for_doc:
            text = _normalize_plain_text_for_extract(text)

    if scan_result is None:
        scan_result = _scan_contract_text(text, file_name=path.name)

    if use_ifilter_for_doc:
        return _build_payload_generic(path, text, scan_result, extract_mode="plain_text_doc")
    return _build_payload_generic(path, text, scan_result, extract_mode="structured_docx")

    text = read_docx(filepath)
    scan_result = scan_docx_for_contract_no(filepath)

    idx_a = text.find("BÊN CHUYỂN NHƯỢNG")
    idx_b = text.find("BÊN NHẬN CHUYỂN NHƯỢNG")
    idx_end = text.find("Hai bên tự nguyện")

    ben_a_section = text[idx_a:idx_b] if idx_a >= 0 and idx_b >= 0 else ""
    ben_b_section = text[idx_b:idx_end] if idx_b >= 0 and idx_end >= 0 else ""

    ben_a = {"nguoi": find_persons(ben_a_section), "dia_chi": find_dia_chi(ben_a_section)}
    ben_b = {"nguoi": find_persons(ben_b_section), "dia_chi": find_dia_chi(ben_b_section)}

    ten_hd = find_ten_hop_dong(text)
    tai_san = find_tai_san(text)
    so_cong_chung = find_so_cong_chung(text) or scan_result.get("contract_no", "")

    return {
        "web_form": {
            "ten_hop_dong": ten_hd,
            "ngay_cong_chung": find_ngay_cong_chung(text),
            "so_cong_chung": so_cong_chung,
            "nhom_hop_dong": guess_nhom_hd(ten_hd),
            "loai_tai_san": guess_loai_tai_san(tai_san, ten_hd),
            "cong_chung_vien": find_ccv(text),
            "thu_ky": DEFAULT_THU_KY,
            "nguoi_yeu_cau": fmt_nguoi_yeu_cau(ben_b),
            "duong_su": fmt_duong_su(ben_a, ben_b),
            "tai_san": tai_san,
        },
        "raw": {
            "ben_a": ben_a,
            "ben_b": ben_b,
            "file_goc": os.path.abspath(filepath),
            "scan_contract_no": scan_result.get("contract_no", ""),
            "scan_reason": scan_result.get("reason", ""),
        },
    }


def main():
    if len(sys.argv) < 2:
        print("Cách dùng:")
        print('  python extract_contract.py "hop_dong.docx"')
        print('  python extract_contract.py "hop_dong.doc"')
        sys.exit(1)

    fp = sys.argv[1]
    if not os.path.exists(fp):
        print(f"Không tìm thấy file: {fp}")
        sys.exit(1)

    result = extract(fp)
    output = json.dumps(result, ensure_ascii=False, indent=2)
    print(output)

    json_path = fp.rsplit(".", 1)[0] + "_extracted.json"
    with open(json_path, "w", encoding="utf-8") as file_obj:
        file_obj.write(output)
    print(f"\n→ Đã lưu: {json_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
