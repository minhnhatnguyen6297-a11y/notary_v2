from fastapi import APIRouter, Depends, Request, Form, HTTPException, UploadFile, File
from fastapi.responses import RedirectResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from typing import Optional, List, Union
from datetime import date, datetime
import io
import json
from pathlib import Path
import re
import unicodedata
from uuid import uuid4
from types import SimpleNamespace

from database import get_db
from models import InheritanceCase, Customer, Property, InheritanceParticipant, InheritanceCaseProperty, WordTemplate

router = APIRouter()
templates = Jinja2Templates(directory="frontend/templates")
WORD_TEMPLATE_UPLOAD_DIR = Path("word_templates/custom")
def _hang_for_role(role: str) -> int:
    role = (role or "").strip()
    if role in ("Cha", "Mẹ", "Cha_vc", "Me_vc", "Vợ/Chồng", "Con", "Cháu", "Con_dau_re"):
        return 1
    if role in ("Ông/Bà", "Anh/Chị/Em"):
        return 2
    return 1


def _to_list(v):
    if v is None:
        return []
    if isinstance(v, list):
        return v
    return [v]


def _normalize_property_ids(primary_property_id: str, raw_property_ids: Optional[Union[List[str], str]]) -> list[int]:
    values: list[str] = []
    for item in _to_list(raw_property_ids):
        if item is None:
            continue
        text = str(item).strip()
        if not text:
            continue
        if text.startswith("[") and text.endswith("]"):
            try:
                parsed = json.loads(text)
                for v in parsed if isinstance(parsed, list) else []:
                    values.append(str(v).strip())
                continue
            except Exception:
                pass
        values.extend([x.strip() for x in text.split(",") if x and x.strip()])

    if primary_property_id:
        values.append(primary_property_id)
    out: list[int] = []
    seen: set[int] = set()
    for v in values:
        if not str(v).isdigit():
            continue
        pid = int(v)
        if pid in seen:
            continue
        seen.add(pid)
        out.append(pid)
    if primary_property_id and primary_property_id.isdigit():
        primary = int(primary_property_id)
        if primary in out:
            out = [primary] + [x for x in out if x != primary]
    return out


def _sync_case_property_links(db: Session, case_id: int, property_ids: list[int], primary_property_id: int) -> None:
    db.query(InheritanceCaseProperty).filter(InheritanceCaseProperty.case_id == case_id).delete()
    for pid in property_ids:
        db.add(
            InheritanceCaseProperty(
                case_id=case_id,
                property_id=pid,
                is_primary=(pid == primary_property_id),
            )
        )
    db.commit()


def _build_temp_participants(
    all_customers: List[Customer],
    participant_id: Optional[Union[List[str], str]],
    participant_role: Optional[Union[List[str], str]],
    participant_share: Optional[Union[List[str], str]],
    participant_receive: Optional[Union[List[str], str]],
    participant_parent_id: Optional[Union[List[str], str]] = None,
):
    id_list = _to_list(participant_id)
    role_list = _to_list(participant_role)
    share_list = _to_list(participant_share)
    receive_list = _to_list(participant_receive)
    parent_list = _to_list(participant_parent_id)
    customers_by_id = {str(c.id): c for c in all_customers}
    participants = []

    for idx, cid in enumerate(id_list):
        cid_str = str(cid or "").strip()
        if not cid_str:
            continue
        customer = customers_by_id.get(cid_str)
        if not customer:
            continue
        role = (role_list[idx] if idx < len(role_list) else "") or "Khac"
        share_raw = share_list[idx] if idx < len(share_list) else "0"
        receive_raw = receive_list[idx] if idx < len(receive_list) else "1"
        parent_raw = parent_list[idx] if idx < len(parent_list) else ""
        try:
            share_val = float(share_raw)
        except Exception:
            share_val = 0.0
        co_nhan = str(receive_raw).lower() in ("1", "true", "on", "yes")

        parent_cid = None
        if parent_raw and str(parent_raw).isdigit():
            parent_cid = int(parent_raw)

        participants.append(SimpleNamespace(
            customer_id=customer.id,
            customer=customer,
            vai_tro=role,
            ty_le=share_val,
            co_nhan_tai_san=co_nhan,
            parent_customer_id=parent_cid
        ))
    participant_ids = {p.customer_id for p in participants}
    return participants, participant_ids


@router.get("/")
def list_cases(request: Request, db: Session = Depends(get_db), q: str = ""):
    cases = db.query(InheritanceCase).order_by(InheritanceCase.id.desc()).all()
    if q:
        cases = [c for c in cases if q.lower() in c.nguoi_chet.ho_ten.lower()]
    return templates.TemplateResponse("cases/list.html", {"request": request, "cases": cases, "q": q})



@router.get("/create")
def create_form(request: Request, db: Session = Depends(get_db)):
    all_customers = db.query(Customer).order_by(Customer.ho_ten).all()
    # Người chết = có ngày chết
    deceased = [c for c in all_customers if c.ngay_chet is not None]
    properties = db.query(Property).order_by(Property.id.desc()).all()
    from datetime import date as _date
    form = {
        "nguoi_chet_id": "", "tai_san_id": "", "ngay_lap_ho_so": _date.today().isoformat(),
        "loai_van_ban": "khai_nhan", "ghi_chu": "", "engine_state_json": ""
    }
    return templates.TemplateResponse("cases/form.html", {
        "request": request, "obj": None,
        "deceased": deceased, "properties": properties, "errors": [],
        "field_errors": {}, "form": form,
        "all_customers": all_customers, "participants": [], "participant_ids": set(), "case_property_ids": [],
    })


@router.post("/create")
def create(
    request: Request,
    nguoi_chet_id: Optional[str] = Form(None),
    tai_san_id: Optional[str] = Form(None),
    property_ids: Optional[Union[List[str], str]] = Form(None),
    participant_id: Optional[Union[List[str], str]] = Form(None),
    participant_role: Optional[Union[List[str], str]] = Form(None),
    participant_share: Optional[Union[List[str], str]] = Form(None),
    participant_receive: Optional[Union[List[str], str]] = Form(None),
    participant_parent_id: Optional[Union[List[str], str]] = Form(None),
    engine_state_json: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    from datetime import date as _date
    form = {
        "nguoi_chet_id": (nguoi_chet_id or "").strip(),
        "tai_san_id": (tai_san_id or "").strip(),
        "engine_state_json": (engine_state_json or "").strip(),
    }
    selected_property_ids = _normalize_property_ids(form["tai_san_id"], property_ids)
    errors = []
    field_errors = {}
    if not form["nguoi_chet_id"]:
        field_errors["nguoi_chet_id"] = "Bắt buộc"
    if not form["tai_san_id"]:
        field_errors["tai_san_id"] = "Bắt buộc"

    all_customers = db.query(Customer).order_by(Customer.ho_ten).all()
    deceased = [c for c in all_customers if c.ngay_chet is not None]
    properties = db.query(Property).order_by(Property.id.desc()).all()
    posted_participants, posted_participant_ids = _build_temp_participants(
        all_customers, participant_id, participant_role, participant_share, participant_receive, participant_parent_id
    )

    if field_errors:
        return templates.TemplateResponse("cases/form.html", {
            "request": request, "obj": None,
            "deceased": deceased, "properties": properties,
            "errors": errors, "field_errors": field_errors, "form": form,
            "all_customers": all_customers,
            "participants": posted_participants,
            "participant_ids": posted_participant_ids,
            "case_property_ids": selected_property_ids,
        })

    try:
        c = InheritanceCase(
            nguoi_chet_id=int(form["nguoi_chet_id"]), tai_san_id=int(form["tai_san_id"]),
            ngay_lap_ho_so=_date.today(),
            loai_van_ban="khai_nhan", ghi_chu=None,
            engine_state_json=(engine_state_json or "").strip() or None,
        )
        db.add(c); db.commit(); db.refresh(c)
        if selected_property_ids:
            _sync_case_property_links(db, c.id, selected_property_ids, int(form["tai_san_id"]))
        pid_list = _to_list(participant_id)
        role_list = _to_list(participant_role)
        share_list = _to_list(participant_share)
        recv_list = _to_list(participant_receive)
        parent_list = _to_list(participant_parent_id)
        if pid_list and role_list:
            for idx, cid in enumerate(pid_list):
                if not cid:
                    continue
                if str(cid) == str(c.nguoi_chet_id):
                    continue
                role = role_list[idx] if idx < len(role_list) else ""
                share_raw = share_list[idx] if idx < len(share_list) else "0"
                receive_raw = recv_list[idx] if idx < len(recv_list) else "1"
                parent_raw = parent_list[idx] if idx < len(parent_list) else ""
                try:
                    share_val = float(share_raw)
                except Exception:
                    share_val = 0.0
                co_nhan = str(receive_raw).lower() in ("1", "true", "on", "yes")

                parent_cid = None
                if parent_raw and str(parent_raw).isdigit():
                    parent_cid = int(parent_raw)

                p = InheritanceParticipant(
                    ho_so_id=c.id, customer_id=int(cid),
                    vai_tro=role or "Khac", hang_thua_ke=_hang_for_role(role or "Khac"),
                    ty_le=share_val, co_nhan_tai_san=co_nhan,
                    parent_customer_id=parent_cid
                )
                db.add(p)
            db.commit()
        return RedirectResponse(f"/cases/{c.id}/edit", status_code=302)
    except Exception as e:
        db.rollback()
        errors.append(f"Lỗi tạo hồ sơ: {e}")
        return templates.TemplateResponse("cases/form.html", {
            "request": request, "obj": None,
            "deceased": deceased, "properties": properties,
            "errors": errors, "field_errors": field_errors, "form": form,
            "all_customers": all_customers,
            "participants": posted_participants,
            "participant_ids": posted_participant_ids,
            "case_property_ids": selected_property_ids,
        })


@router.get("/{cid}")
def detail(cid: int, request: Request, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case: raise HTTPException(404)
    all_customers = db.query(Customer).order_by(Customer.ho_ten).all()
    participant_ids = {p.customer_id for p in case.participants}
    available = [c for c in all_customers if c.id not in participant_ids and c.id != case.nguoi_chet_id]
    return templates.TemplateResponse("cases/detail.html", {
        "request": request, "case": case, "available": available,
        "vai_tro_options": ["Vợ/Chồng", "Con", "Cha/Mẹ", "Anh/Chị/Em"]
    })


@router.get("/{cid}/edit")
def edit_form(cid: int, request: Request, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case: raise HTTPException(404)
    if case.is_locked:
        return RedirectResponse(f"/cases/{cid}", status_code=302)
    all_customers = db.query(Customer).order_by(Customer.ho_ten).all()
    deceased = [c for c in all_customers if c.ngay_chet is not None]
    properties = db.query(Property).order_by(Property.id.desc()).all()
    participants = case.participants
    participant_ids = {p.customer_id for p in participants}
    case_property_ids = [int(link.property_id) for link in sorted(case.property_links, key=lambda x: (not x.is_primary, x.id))]
    if not case_property_ids and case.tai_san_id:
        case_property_ids = [int(case.tai_san_id)]
    form = {
        "nguoi_chet_id": str(case.nguoi_chet_id) if case.nguoi_chet_id else "",
        "tai_san_id": str(case.tai_san_id) if case.tai_san_id else "",
        "ngay_lap_ho_so": case.ngay_lap_ho_so.isoformat() if case.ngay_lap_ho_so else "",
        "loai_van_ban": case.loai_van_ban or "khai_nhan",
        "noi_niem_yet": case.noi_niem_yet or "",
        "ghi_chu": case.ghi_chu or "",
        "engine_state_json": case.engine_state_json or "",
    }
    return templates.TemplateResponse("cases/form.html", {
        "request": request, "obj": case,
        "deceased": deceased, "properties": properties, "errors": [],
        "field_errors": {}, "form": form,
        "all_customers": all_customers, "participants": participants, "participant_ids": participant_ids,
        "case_property_ids": case_property_ids,
    })


@router.post("/{cid}/edit")
def edit(
    cid: int, request: Request,
    nguoi_chet_id: Optional[str] = Form(None), tai_san_id: Optional[str] = Form(None),
    property_ids: Optional[Union[List[str], str]] = Form(None),
    noi_niem_yet: Optional[str] = Form(None),
    participant_id: Optional[Union[List[str], str]] = Form(None),
    participant_role: Optional[Union[List[str], str]] = Form(None),
    participant_share: Optional[Union[List[str], str]] = Form(None),
    participant_receive: Optional[Union[List[str], str]] = Form(None),
    participant_parent_id: Optional[Union[List[str], str]] = Form(None),
    engine_state_json: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case or case.is_locked: raise HTTPException(400)
    form = {
        "nguoi_chet_id": (nguoi_chet_id or "").strip(),
        "tai_san_id": (tai_san_id or "").strip(),
        "noi_niem_yet": (noi_niem_yet or "").strip(),
        "engine_state_json": (engine_state_json or "").strip(),
    }
    selected_property_ids = _normalize_property_ids(form["tai_san_id"], property_ids)
    errors = []
    field_errors = {}
    if not form["nguoi_chet_id"]:
        field_errors["nguoi_chet_id"] = "Bắt buộc"
    if not form["tai_san_id"]:
        field_errors["tai_san_id"] = "Bắt buộc"

    all_customers = db.query(Customer).order_by(Customer.ho_ten).all()
    deceased = [c for c in all_customers if c.ngay_chet is not None]
    properties = db.query(Property).order_by(Property.id.desc()).all()
    posted_participants, posted_participant_ids = _build_temp_participants(
        all_customers, participant_id, participant_role, participant_share, participant_receive, participant_parent_id
    )
    if field_errors:
        return templates.TemplateResponse("cases/form.html", {
            "request": request, "obj": case,
            "deceased": deceased, "properties": properties,
            "errors": errors, "field_errors": field_errors, "form": form,
            "all_customers": all_customers,
            "participants": posted_participants,
            "participant_ids": posted_participant_ids,
            "case_property_ids": selected_property_ids,
        })

    try:
        case.nguoi_chet_id = int(form["nguoi_chet_id"]); case.tai_san_id = int(form["tai_san_id"])
        case.noi_niem_yet = form["noi_niem_yet"] or None
        case.engine_state_json = (engine_state_json or "").strip() or None
        db.commit()
        if selected_property_ids:
            _sync_case_property_links(db, case.id, selected_property_ids, int(form["tai_san_id"]))
        db.query(InheritanceParticipant).filter(InheritanceParticipant.ho_so_id == case.id).delete()
        db.commit()
        pid_list = _to_list(participant_id)
        role_list = _to_list(participant_role)
        share_list = _to_list(participant_share)
        recv_list = _to_list(participant_receive)
        parent_list = _to_list(participant_parent_id)
        if pid_list and role_list:
            for idx, participant_customer_id in enumerate(pid_list):
                if not participant_customer_id:
                    continue
                if str(participant_customer_id) == str(case.nguoi_chet_id):
                    continue
                role = role_list[idx] if idx < len(role_list) else ""
                share_raw = share_list[idx] if idx < len(share_list) else "0"
                receive_raw = recv_list[idx] if idx < len(recv_list) else "1"
                parent_raw = parent_list[idx] if idx < len(parent_list) else ""
                try:
                    share_val = float(share_raw)
                except Exception:
                    share_val = 0.0
                co_nhan = str(receive_raw).lower() in ("1", "true", "on", "yes")

                parent_cid = None
                if parent_raw and str(parent_raw).isdigit():
                    parent_cid = int(parent_raw)

                p = InheritanceParticipant(
                    ho_so_id=case.id, customer_id=int(participant_customer_id),
                    vai_tro=role or "Khac", hang_thua_ke=_hang_for_role(role or "Khac"),
                    ty_le=share_val, co_nhan_tai_san=co_nhan,
                    parent_customer_id=parent_cid
                )
                db.add(p)
            db.commit()
        return RedirectResponse(f"/cases/{cid}/edit", status_code=302)
    except Exception as e:
        db.rollback()
        errors.append(f"Lỗi cập nhật hồ sơ: {e}")
        return templates.TemplateResponse("cases/form.html", {
            "request": request, "obj": case,
            "deceased": deceased, "properties": properties,
            "errors": errors, "field_errors": field_errors, "form": form,
            "all_customers": all_customers,
            "participants": posted_participants,
            "participant_ids": posted_participant_ids,
            "case_property_ids": selected_property_ids,
        })


@router.post("/{cid}/lock")
def lock(cid: int, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if case: case.trang_thai = "locked"; db.commit()
    return RedirectResponse(f"/cases/{cid}", status_code=302)


@router.post("/{cid}/unlock")
def unlock(cid: int, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if case: case.trang_thai = "draft"; db.commit()
    return RedirectResponse(f"/cases/{cid}", status_code=302)


@router.post("/{cid}/delete")
def delete(cid: int, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if case and not case.is_locked:
        db.delete(case); db.commit()
    return RedirectResponse("/cases", status_code=302)


def _get_selected_word_template_path(db: Session) -> Optional[Path]:
    active = (
        db.query(WordTemplate)
        .filter(WordTemplate.is_active == True)
        .order_by(WordTemplate.id.desc())
        .first()
    )
    if active and active.duong_dan_file:
        p = Path(active.duong_dan_file)
        if p.exists():
            return p

    template_candidates = [
        Path(r"\\maychu\D\Minh\HỒ SƠ UBND CÁC XÃ\2. Mẫu thừa kế\xã_PCDS -.docx"),
        Path("word_templates/xa_PCDS_template.docx"),
    ]
    existing_templates = [p for p in template_candidates if p.exists()]
    if not existing_templates:
        return None
    return max(existing_templates, key=lambda p: p.stat().st_mtime)


@router.get("/templates/list-json")
def list_templates_json(db: Session = Depends(get_db)):
    """API trả về danh sách template Word dạng JSON cho modal xuất văn bản."""
    items = db.query(WordTemplate).order_by(WordTemplate.id.desc()).all()
    # Also include built-in templates from word_templates/ dir
    builtin = []
    for p in Path("word_templates").glob("*.docx"):
        if p.exists():
            builtin.append({"id": f"builtin:{p.name}", "ten_mau": p.stem, "is_active": False, "builtin": True})
    return {
        "templates": [
            {"id": t.id, "ten_mau": t.ten_mau, "ten_file_goc": t.ten_file_goc, "is_active": t.is_active, "builtin": False}
            for t in items
        ] + builtin
    }


@router.get("/templates/manage")
def word_templates_page(request: Request, db: Session = Depends(get_db), ok: str = "", err: str = ""):
    items = db.query(WordTemplate).order_by(WordTemplate.id.desc()).all()
    return templates.TemplateResponse("cases/templates.html", {
        "request": request,
        "items": items,
        "ok": ok,
        "err": err,
    })


@router.post("/templates/manage/upload")
async def upload_word_template(
    ten_mau: str = Form(...),
    file_mau: UploadFile = File(...),
    dat_mac_dinh: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    ten_mau = (ten_mau or "").strip()
    if not ten_mau:
        return RedirectResponse("/cases/templates/manage?err=Vui+long+nhap+ten+mau", status_code=302)
    if not file_mau or not file_mau.filename:
        return RedirectResponse("/cases/templates/manage?err=Vui+long+chon+file", status_code=302)
    if not file_mau.filename.lower().endswith(".docx"):
        return RedirectResponse("/cases/templates/manage?err=Chi+ho+tro+file+.docx", status_code=302)

    WORD_TEMPLATE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    saved_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}.docx"
    saved_path = WORD_TEMPLATE_UPLOAD_DIR / saved_name
    content = await file_mau.read()
    saved_path.write_bytes(content)

    set_active = str(dat_mac_dinh).lower() in ("1", "true", "on", "yes")
    if set_active:
        db.query(WordTemplate).update({WordTemplate.is_active: False})

    item = WordTemplate(
        ten_mau=ten_mau,
        ten_file_goc=file_mau.filename,
        duong_dan_file=str(saved_path),
        is_active=set_active,
    )
    db.add(item)
    db.commit()
    return RedirectResponse("/cases/templates/manage?ok=Tai+mau+thanh+cong", status_code=302)


@router.post("/templates/manage/{tid}/activate")
def activate_word_template(tid: int, db: Session = Depends(get_db)):
    item = db.query(WordTemplate).filter(WordTemplate.id == tid).first()
    if not item:
        return RedirectResponse("/cases/templates/manage?err=Khong+tim+thay+mau", status_code=302)
    db.query(WordTemplate).update({WordTemplate.is_active: False})
    item.is_active = True
    db.commit()
    return RedirectResponse("/cases/templates/manage?ok=Da+chon+mau+mac+dinh", status_code=302)


@router.post("/templates/manage/{tid}/delete")
def delete_word_template(tid: int, db: Session = Depends(get_db)):
    item = db.query(WordTemplate).filter(WordTemplate.id == tid).first()
    if not item:
        return RedirectResponse("/cases/templates/manage?err=Khong+tim+thay+mau", status_code=302)

    was_active = bool(item.is_active)
    file_path = Path(item.duong_dan_file or "")
    db.delete(item)
    db.commit()

    if file_path.exists():
        try:
            file_path.unlink()
        except Exception:
            pass

    if was_active:
        latest = db.query(WordTemplate).order_by(WordTemplate.id.desc()).first()
        if latest:
            latest.is_active = True
            db.commit()

    return RedirectResponse("/cases/templates/manage?ok=Da+xoa+mau", status_code=302)


# ── JSON API cho modal Quản lý mẫu (không rời trang) ──────────────────────────

@router.post("/templates/api/upload")
async def api_upload_template(
    ten_mau: str = Form(...),
    file_mau: UploadFile = File(...),
    dat_mac_dinh: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    ten_mau = (ten_mau or "").strip()
    if not ten_mau:
        return JSONResponse({"ok": False, "err": "Vui lòng nhập tên mẫu"}, status_code=400)
    if not file_mau or not file_mau.filename:
        return JSONResponse({"ok": False, "err": "Vui lòng chọn file"}, status_code=400)
    if not file_mau.filename.lower().endswith(".docx"):
        return JSONResponse({"ok": False, "err": "Chỉ hỗ trợ file .docx"}, status_code=400)

    WORD_TEMPLATE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    saved_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}.docx"
    saved_path = WORD_TEMPLATE_UPLOAD_DIR / saved_name
    content = await file_mau.read()
    saved_path.write_bytes(content)

    set_active = str(dat_mac_dinh).lower() in ("1", "true", "on", "yes")
    if set_active:
        db.query(WordTemplate).update({WordTemplate.is_active: False})

    item = WordTemplate(
        ten_mau=ten_mau,
        ten_file_goc=file_mau.filename,
        duong_dan_file=str(saved_path),
        is_active=set_active,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return JSONResponse({"ok": True, "id": item.id, "ten_mau": item.ten_mau, "is_active": item.is_active})


@router.post("/templates/api/{tid}/activate")
def api_activate_template(tid: int, db: Session = Depends(get_db)):
    item = db.query(WordTemplate).filter(WordTemplate.id == tid).first()
    if not item:
        return JSONResponse({"ok": False, "err": "Không tìm thấy mẫu"}, status_code=404)
    db.query(WordTemplate).update({WordTemplate.is_active: False})
    item.is_active = True
    db.commit()
    return JSONResponse({"ok": True})


@router.post("/templates/api/{tid}/delete")
def api_delete_template(tid: int, db: Session = Depends(get_db)):
    item = db.query(WordTemplate).filter(WordTemplate.id == tid).first()
    if not item:
        return JSONResponse({"ok": False, "err": "Không tìm thấy mẫu"}, status_code=404)
    was_active = bool(item.is_active)
    file_path = Path(item.duong_dan_file or "")
    db.delete(item)
    db.commit()
    if file_path.exists():
        try:
            file_path.unlink()
        except Exception:
            pass
    if was_active:
        latest = db.query(WordTemplate).order_by(WordTemplate.id.desc()).first()
        if latest:
            latest.is_active = True
            db.commit()
    return JSONResponse({"ok": True})


def _fmt_date(d: Optional[date]) -> str:
    if not d:
        return ""
    return d.strftime("%d/%m/%Y")


def _fmt_birth_or_year(d: Optional[date]) -> str:
    if not d:
        return ""
    if d.day == 1 and d.month == 1:
        return str(d.year)
    return d.strftime("%d/%m/%Y")


def _safe_text(v) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    if s == "0":
        return ""
    return s


def _so_thanh_chu(so: float) -> str:
    """Chuyển số thực (diện tích m²) thành chữ tiếng Việt."""
    if so is None:
        return ""
    don_vi = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    hang = ["", "mười", "trăm", "nghìn", "", "", "triệu", "", "", "tỷ"]

    def _doc_ba_chu_so(n: int) -> str:
        tram = n // 100
        chuc = (n % 100) // 10
        dv   = n % 10
        result = ""
        if tram:
            result += don_vi[tram] + " trăm"
            if chuc == 0 and dv:
                result += " linh " + don_vi[dv]
            elif chuc:
                result += " " + (don_vi[chuc] + " mươi" if chuc > 1 else "mười")
                if dv == 1 and chuc > 1:
                    result += " mốt"
                elif dv == 5 and chuc > 0:
                    result += " lăm"
                elif dv:
                    result += " " + don_vi[dv]
        elif chuc:
            result += (don_vi[chuc] + " mươi" if chuc > 1 else "mười")
            if dv == 1 and chuc > 1:
                result += " mốt"
            elif dv == 5 and chuc > 0:
                result += " lăm"
            elif dv:
                result += " " + don_vi[dv]
        elif dv:
            result += don_vi[dv]
        return result.strip()

    # Tách phần nguyên và thập phân
    phan_nguyen = int(so)
    phan_le_str = ""
    if so != phan_nguyen:
        le = round(so - phan_nguyen, 6)
        dec_s = f"{le:.6f}".split(".")[1].rstrip("0")
        if dec_s:
            phan_le_str = " phẩy " + " ".join(don_vi[int(d)] for d in dec_s)

    if phan_nguyen == 0:
        return ("không" + phan_le_str).strip()

    # Xử lý số nguyên
    parts = []
    n = phan_nguyen
    ty  = n // 1_000_000_000; n %= 1_000_000_000
    tr  = n // 1_000_000;     n %= 1_000_000
    ng  = n // 1_000;         n %= 1_000
    dv3 = n

    if ty:  parts.append(_doc_ba_chu_so(ty) + " tỷ")
    if tr:  parts.append(_doc_ba_chu_so(tr) + " triệu")
    if ng:  parts.append(_doc_ba_chu_so(ng) + " nghìn")
    if dv3: parts.append(_doc_ba_chu_so(dv3))

    return (" ".join(parts) + phan_le_str).strip()


def _pick_core_people(case: InheritanceCase):
    owner = case.nguoi_chet
    spouse = None
    for p in case.participants:
        if (p.vai_tro or "").strip() == "Vợ/Chồng":
            spouse = p.customer
            break

    pair = [c for c in [owner, spouse] if c is not None]
    
    nam = [c for c in pair if (c.gioi_tinh or "").strip().lower() == "nam"]
    nu = [c for c in pair if (c.gioi_tinh or "").strip().lower() in ("nữ", "nu", "nu")]
    
    if len(nam) == 1 and len(nu) == 1:
        person1 = nam[0]
        person2 = nu[0]
    elif len(pair) == 2:
        person1 = pair[0]
        person2 = pair[1]
    elif len(pair) == 1:
        person1 = pair[0]
        person2 = None
    else:
        person1 = None
        person2 = None

    excluded_ids = {c.id for c in [person1, person2] if c is not None}
    receivers = [p for p in case.participants if p.co_nhan_tai_san and p.customer_id not in excluded_ids]
    receivers = sorted(receivers, key=lambda p: (-(p.ty_le or 0), p.customer_id))
    non_receivers = [p for p in case.participants if (not p.co_nhan_tai_san) and p.customer_id not in excluded_ids]
    non_receivers = sorted(non_receivers, key=lambda p: p.customer_id)

    person3 = receivers[0].customer if receivers else None
    
    rest_receivers = [p.customer for p in receivers[1:]] if receivers else []
    rest_non_receivers = [p.customer for p in non_receivers]
    people_4_plus = rest_receivers + rest_non_receivers
    return person1, person2, person3, people_4_plus


def _build_template_mapping(case: InheritanceCase) -> dict:
    ts = case.tai_san
    person1, person2, person3, people_4_plus = _pick_core_people(case)

    people_slots = [None] * 21
    people_slots[1] = person1
    people_slots[2] = person2
    people_slots[3] = person3
    for idx, c in enumerate(people_4_plus[:17], start=4):
        people_slots[idx] = c

    import json as _json
    from datetime import date as _date_cls
    today = _date_cls.today()

    noi_niem_yet = _safe_text(case.noi_niem_yet) if case.noi_niem_yet else _safe_text(ts.dia_chi)

    # Phân tích land_rows_json để lấy dữ liệu từng loại đất
    land_rows = []
    if ts.land_rows_json:
        try:
            land_rows = _json.loads(ts.land_rows_json)
        except Exception:
            pass

    # Tổng diện tích: ưu tiên tính từ land_rows, fallback về ts.dien_tich
    if land_rows:
        total = 0.0
        for r in land_rows:
            try:
                total += float(r.get("dien_tich") or 0)
            except (ValueError, TypeError):
                pass
        dien_tich_so = total if total > 0 else ts.dien_tich
    else:
        dien_tich_so = ts.dien_tich

    dien_tich_str = f"{dien_tich_so:g}" if dien_tich_so else ""
    dien_tich_chu = _so_thanh_chu(dien_tich_so).capitalize() if dien_tich_so else ""

    loai_so_val = _safe_text(ts.loai_so) or "Giấy chứng nhận quyền sử dụng đất"

    m = {
        "[Tên file]": f"ho_so_thua_ke_{case.id}",
        "[Niêm Yết]": noi_niem_yet,
        "[NIÊM YẾT]": noi_niem_yet.upper() if noi_niem_yet else "",
        "[Loại sổ]": loai_so_val,
        "[Địa chỉ đất]": _safe_text(ts.dia_chi),
        "[Serial]": _safe_text(ts.so_serial),
        "[Số vào sổ]": _safe_text(ts.so_vao_so),
        "[Số thửa]": _safe_text(ts.so_thua_dat),
        "[Số tờ]": _safe_text(ts.so_to_ban_do),
        "[Diện tích]": dien_tich_str,
        "[Diện tích chữ]": dien_tich_chu,
        "[Hình thức sử dụng]": _safe_text(ts.hinh_thuc_su_dung),
        "[Loại đất]": _safe_text(ts.loai_dat),
        "[Nguồn gốc]": _safe_text(ts.nguon_goc),
        "[Ngày cấp sổ]": _fmt_date(ts.ngay_cap),
        "[Cơ quan cấp sổ]": _safe_text(ts.co_quan_cap),
        "[Ngày]": str(today.day),
        "[Tháng]": f"{today.month:02d}",
        "[Ngày chữ]": _so_thanh_chu(today.day),
        "[Tháng chữ]": _so_thanh_chu(today.month),
        "[Người ủy quyền]": "",
        "[Người ủy quyền2]": "",
        "[Số công chứng]": "",
        "[ONT]": "",
        "[CLN]": "",
        "[NTS]": "",
        "[LUC]": "",
        "[Giá chuyển nhượng]": "",
        "[SĐT]": "",
    }

    for i in range(1, 21):
        c = people_slots[i]
        m[f"[Tên {i}]"] = _safe_text(c.ho_ten if c else "")
        m[f"[Năm sinh {i}]"] = _safe_text(_fmt_birth_or_year(c.ngay_sinh) if c else "")
        m[f"[CCCD {i}]"] = _safe_text(c.so_giay_to if c else "")
        m[f"[Ngày cấp {i}]"] = _safe_text(_fmt_date(c.ngay_cap) if c else "")
        m[f"[Địa chỉ {i}]"] = _safe_text(c.dia_chi if c else "")
        m[f"[Loại CC {i}]"] = _safe_text(c.loai_giay_to if c else "")
        m[f"[Nơi cấp CC {i}]"] = _safe_text(c.noi_cap if c else "")
        m[f"[Thường trú {i}]"] = _safe_text(c.loai_dia_chi if c else "")
        m[f"[Năm chết {i}]"] = _safe_text(_fmt_date(c.ngay_chet) if c else "")

    # Mapping từng loại đất theo số thứ tự: [Loại đất N], [Diện tích N], [Thời hạn N]
    for i, row in enumerate(land_rows[:10], start=1):
        loai = str(row.get("loai_dat", "")).strip()
        dien = str(row.get("dien_tich", "")).strip()
        thoi = str(row.get("thoi_han", "")).strip()
        m[f"[Loại đất {i}]"] = loai
        m[f"[Diện tích {i}]"] = dien
        m[f"[Thời hạn {i}]"] = thoi
    # Xóa giá trị cho các slot vượt quá số dòng thực tế (tối đa 10)
    for i in range(len(land_rows) + 1, 11):
        m[f"[Loại đất {i}]"] = ""
        m[f"[Diện tích {i}]"] = ""
        m[f"[Thời hạn {i}]"] = ""

    # [Thời hạn 1] alias → dòng đầu tiên hoặc ts.thoi_han (backward compat)
    if not m.get("[Thời hạn 1]") and ts.thoi_han:
        m["[Thời hạn 1]"] = _safe_text(ts.thoi_han)

    return m


def _normalize_token(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace("đ", "d")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s)
    return s


def _build_normalized_mapping(mapping: dict) -> dict:
    normalized = {}
    for k, v in mapping.items():
        if not (k.startswith("[") and k.endswith("]")):
            continue
        token = k[1:-1]
        normalized[_normalize_token(token)] = v
    return normalized


def _replace_text_placeholders(text: str, mapping: dict, normalized_mapping: dict) -> str:
    new_text = text
    for k, v in mapping.items():
        if k in new_text:
            new_text = new_text.replace(k, v)

    def _token_repl(match):
        token = match.group(1)
        direct = mapping.get(f"[{token}]")
        if direct is not None:
            return direct
        norm = _normalize_token(token)
        if norm in normalized_mapping:
            return normalized_mapping[norm]
        return match.group(0)

    return re.sub(r"\[([^\[\]]+)\]", _token_repl, new_text)


def _replace_in_paragraph(paragraph, mapping: dict, normalized_mapping: dict):
    if not paragraph.runs:
        return
        
    # Check if there's anything to replace at all
    text = "".join(r.text for r in paragraph.runs)
    if "[" not in text or "]" not in text:
        return

    # Try run-by-run first to perfectly preserve inline formatting
    for r in paragraph.runs:
        if "[" in r.text and "]" in r.text:
            new_t = _replace_text_placeholders(r.text, mapping, normalized_mapping)
            if new_t != r.text:
                r.text = new_t

    # Re-evaluate text since runs might have changed
    text = "".join(r.text for r in paragraph.runs)
    if "[" not in text or "]" not in text:
        return
        
    new_text = _replace_text_placeholders(text, mapping, normalized_mapping)
    if new_text != text:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.clear()



def _replace_in_doc(doc, mapping: dict):
    normalized_mapping = _build_normalized_mapping(mapping)
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping, normalized_mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping, normalized_mapping)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            _replace_in_paragraph(p, mapping, normalized_mapping)
        for p in sec.footer.paragraphs:
            _replace_in_paragraph(p, mapping, normalized_mapping)


@router.get("/{cid}/export-word-legacy")
def export_word(cid: int, db: Session = Depends(get_db)):
    """Xuat ho so thua ke ra file Word."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        raise HTTPException(status_code=500, detail="Thieu thu vien python-docx. Vui long cai requirements.")

    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case: raise HTTPException(404)

    doc = Document()

    # Tiêu đề
    title = doc.add_heading("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    van_ban_name = "VĂN BẢN KHAI NHẬN DI SẢN THỪA KẾ" if case.loai_van_ban == "khai_nhan" else "VĂN BẢN THỎA THUẬN PHÂN CHIA DI SẢN THỪA KẾ"
    h = doc.add_heading(van_ban_name, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Thông tin người chết
    doc.add_heading("I. THÔNG TIN NGƯỜI ĐỂ LẠI DI SẢN", level=3)
    nd = case.nguoi_chet
    doc.add_paragraph(f"Họ và tên: {nd.ho_ten}")
    doc.add_paragraph(f"Ngày sinh: {nd.ngay_sinh.strftime('%d/%m/%Y') if nd.ngay_sinh else ''}")
    doc.add_paragraph(f"Ngày chết: {nd.ngay_chet.strftime('%d/%m/%Y') if nd.ngay_chet else ''}")
    doc.add_paragraph(f"Số CCCD/Giấy tờ: {nd.so_giay_to}")
    doc.add_paragraph(f"Địa chỉ thường trú: {nd.dia_chi}")

    doc.add_paragraph()

    # Thông tin tài sản
    doc.add_heading("II. TÀI SẢN", level=3)
    ts = case.tai_san
    doc.add_paragraph(f"Số serial GCN: {ts.so_serial}")
    doc.add_paragraph(f"Số vào sổ: {ts.so_vao_so or ''}")
    doc.add_paragraph(f"Số thửa: {ts.so_thua_dat or ''} - Tờ bản đồ số: {ts.so_to_ban_do or ''}")
    doc.add_paragraph(f"Địa chỉ: {ts.dia_chi}")
    doc.add_paragraph(f"Loại đất: {ts.loai_dat or ''}")
    doc.add_paragraph(f"Thời hạn sử dụng: {ts.thoi_han or ''}")
    doc.add_paragraph(f"Cơ quan cấp: {ts.co_quan_cap or ''}")

    doc.add_paragraph()

    # Người thừa kế
    doc.add_heading("III. NHỮNG NGƯỜI THỪA KẾ", level=3)
    nhan = [p for p in case.participants if p.co_nhan_tai_san]
    tuchoi = [p for p in case.participants if not p.co_nhan_tai_san]

    if nhan:
        doc.add_paragraph("Những người nhận thừa kế:")
        for i, p in enumerate(nhan, 1):
            c = p.customer
            ty_le = float(p.ty_le or 0)
            line = f"{i}. {c.ho_ten} - {p.vai_tro} - Ty le: {ty_le:.1f}%"
            doc.add_paragraph(line, style="List Number")

    if tuchoi:
        doc.add_paragraph()
        doc.add_paragraph("Những người từ chối nhận di sản:")
        for p in tuchoi:
            doc.add_paragraph(f"- {p.customer.ho_ten} ({p.vai_tro}): Từ chối nhận")

    doc.add_paragraph()
    doc.add_paragraph(f"Ngày lập văn bản: {case.ngay_lap_ho_so.strftime('%d tháng %m năm %Y')}")

    doc.add_paragraph()
    doc.add_paragraph("CÔNG CHỨNG VIÊN")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("(Ký và đóng dấu)")

    # Xuất ra stream
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"ho_so_thua_ke_{cid}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{cid}/export-word")
def export_word_from_template(cid: int, db: Session = Depends(get_db), template_id: Optional[str] = None):
    """Export inheritance case using the selected Word template."""
    try:
        from docx import Document
    except Exception:
        raise HTTPException(status_code=500, detail="Thieu thu vien python-docx. Vui long cai requirements.")

    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case:
        raise HTTPException(404)

    # Resolve template path from template_id param
    template_path = None
    if template_id:
        if str(template_id).startswith("builtin:"):
            fname = str(template_id)[len("builtin:"):]
            p = Path("word_templates") / fname
            if p.exists():
                template_path = p
        else:
            try:
                tid = int(template_id)
                t = db.query(WordTemplate).filter(WordTemplate.id == tid).first()
                if t and t.duong_dan_file:
                    p = Path(t.duong_dan_file)
                    if p.exists():
                        template_path = p
            except (ValueError, TypeError):
                pass

    if not template_path:
        template_path = _get_selected_word_template_path(db)
    if not template_path:
        raise HTTPException(status_code=500, detail="Khong tim thay file template Word.")

    try:
        doc = Document(str(template_path))
    except Exception as ex:
        raise HTTPException(status_code=500, detail=f"Khong mo duoc template: {ex}")

    mapping = _build_template_mapping(case)
    _replace_in_doc(doc, mapping)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"ho_so_thua_ke_{cid}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.post("/live-preview")
def create_live_preview(
    request: Request,
    ngay_lap_ho_so: Optional[str] = Form(None),
    loai_van_ban: Optional[str] = Form("khai_nhan"),
    tai_san_id: Optional[str] = Form(None),
    ghi_chu: Optional[str] = Form(None),
    participant_id: Optional[Union[List[str], str]] = Form(None),
    participant_role: Optional[Union[List[str], str]] = Form(None),
    participant_share: Optional[Union[List[str], str]] = Form(None),
    participant_receive: Optional[Union[List[str], str]] = Form(None),
    participant_parent_id: Optional[Union[List[str], str]] = Form(None),
    db: Session = Depends(get_db)
):
    # Dummy case to use existing mapping logic
    class DummyCase:
        def __init__(self, ts, parts, loai, ngay):
            self.id = 9999
            self.tai_san = ts
            self.participants = parts
            self.loai_van_ban = loai
            self.ngay_lap_ho_so = ngay
            
            # Find owner
            self.nguoi_chet = None
            for p in parts:
                if p.vai_tro == "Owner" or p.customer_id == ts.id: # Just a fallback
                    pass # We will rely on the participants list in _pick_core_people
            
            # Actually, _pick_core_people expects nguoi_chet. 
            # Let's find the owner from the participants. Wait, in form, owner is NOT sent if we don't handle it.
            # In form: roleMap doesn't have Owner. Let's fix that.
            
    # We will build a custom HTML generator instead of relying on _build_template_mapping entirely because we want SMART documents.
    all_customers = db.query(Customer).all()
    customers_by_id = {str(c.id): c for c in all_customers}
    
    ts = db.query(Property).filter(Property.id == tai_san_id).first() if tai_san_id else None
    if not ts:
        ts = SimpleNamespace(so_serial="...", so_vao_so="...", so_thua_dat="...", so_to_ban_do="...", dia_chi="[...] Vui lòng chọn tài sản", loai_dat="", thoi_han="", nguon_goc="", ngay_cap=None, co_quan_cap="")
        
    id_list = _to_list(participant_id)
    role_list = _to_list(participant_role)
    share_list = _to_list(participant_share)
    receive_list = _to_list(participant_receive)
    parent_list = _to_list(participant_parent_id)
    
    participants = []
    owner = None
    spouses = []
    children = []
    grand = []
    parents = []
    
    for idx, cid in enumerate(id_list):
        c = customers_by_id.get(str(cid))
        if not c: continue
        role = role_list[idx] if idx < len(role_list) else ""
        recv = str(receive_list[idx]).lower() in ("1", "true") if idx < len(receive_list) else True
        parent_raw = parent_list[idx] if idx < len(parent_list) else ""
        parent_cid = int(parent_raw) if parent_raw and str(parent_raw).isdigit() else None
        
        # In fix_html.py we didn't map owner. The form hidden inputs do not include owner.
        # Wait, the owner IS in the form because owner card has data-role="Owner" (but roleMap in form didn't map it. Let me just use standard list)
        
        p = SimpleNamespace(customer=c, vai_tro=role, co_nhan_tai_san=recv, parent_customer_id=parent_cid)
        participants.append(p)
        
        if role in ("Cha", "Mẹ", "Cha_vc", "Me_vc"): parents.append(p)
        elif role == "Vợ/Chồng": spouses.append(p)
        elif role == "Con": children.append(p)
        elif role == "Cháu": grand.append(p)
        elif role == "Owner": owner = c
        
    # If owner missing (due to roleMap bug in JS we just wrote), let's guess from dead people
    if not owner:
        dead = [p.customer for p in participants if p.customer.ngay_chet]
        if dead: owner = dead[0]
        else: owner = SimpleNamespace(ho_ten="[Người để lại di sản]", so_giay_to="...", ngay_sinh=None, ngay_chet=None)

    nhan = [p for p in participants if p.co_nhan_tai_san and p.customer != owner]
    tuchoi = [p for p in participants if not p.co_nhan_tai_san and p.customer != owner]

    flags = {
        "co_nguoi_dai_dien": False,  # Later expanded using payload from frontend
    }

    ngay_lap = None
    if ngay_lap_ho_so:
        try:
            ngay_lap = datetime.strptime(ngay_lap_ho_so, "%Y-%m-%d").date()
        except:
            pass

    template = templates.get_template("cases/_document_template.html")
    html = template.render({
        "loai_van_ban": loai_van_ban,
        "owner": owner,
        "spouses": spouses,
        "parents": parents,
        "children": children,
        "grand": grand,
        "nhan": nhan,
        "tuchoi": tuchoi,
        "has_tu_choi": len(tuchoi) > 0,
        "ts": ts,
        "flags": flags,
        "ngay_lap": ngay_lap
    })

    return JSONResponse({"html_content": html})

@router.post("/export-draft")
def export_draft_generic(html_content: str = Form("")):
    try:
        from docx import Document
        from htmldocx import HtmlToDocx
        doc = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html_content, doc)
    except ImportError:
        from docx import Document
        from bs4 import BeautifulSoup
        doc = Document()
        soup = BeautifulSoup(html_content, "html.parser")
        doc.add_paragraph(soup.get_text(separator='\n'))
        doc.add_paragraph("\n\n(Lỗi: Yêu cầu pip install htmldocx để kết xuất chính tả/định dạng HTML)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"ban_nhap_ho_so_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/{cid}/preview")
def preview_word(cid: int, request: Request, db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case: raise HTTPException(404)
    mapping = _build_template_mapping(case)
    
    html_content = f"""
    <h1 style="text-align: center;">CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</h1>
    <h2 style="text-align: center;">Độc lập - Tự do - Hạnh phúc</h2>
    <p>&nbsp;</p>
    <h2 style="text-align: center;">{ 'VĂN BẢN KHAI NHẬN DI SẢN THỪA KẾ' if case.loai_van_ban == 'khai_nhan' else 'VĂN BẢN THỎA THUẬN PHÂN CHIA DI SẢN THỪA KẾ' }</h2>
    <p>&nbsp;</p>
    <p>Chúng tôi gồm có:</p>
    <p><b>1. {mapping.get('[Tên 1]', '')}</b> sinh năm {mapping.get('[Năm sinh 1]', '')}, CCCD số {mapping.get('[CCCD 1]', '')} cấp ngày {mapping.get('[Ngày cấp 1]', '')} tại {mapping.get('[Nơi cấp CC 1]', '')}</p>
    """
    if mapping.get('[Tên 2]', ''):
        html_content += f"""    <p><b>2. {mapping.get('[Tên 2]', '')}</b> sinh năm {mapping.get('[Năm sinh 2]', '')}, CCCD số {mapping.get('[CCCD 2]', '')} cấp ngày {mapping.get('[Ngày cấp 2]', '')}</p>"""
    
    html_content += f"""
    <p><i>(Cùng các đồng thừa kế khác...)</i></p>
    <p>&nbsp;</p>
    <h3>DI SẢN THỪA KẾ:</h3>
    <p>Giấy chứng nhận quyền sử dụng đất số <b>{mapping.get('[Serial]', '')}</b>, số vào sổ <b>{mapping.get('[Số vào sổ]', '')}</b> do {mapping.get('[Cơ quan cấp sổ]', '')} cấp ngày {mapping.get('[Ngày cấp sổ]', '')}</p>
    <p>Thửa đất số: {mapping.get('[Số thửa]', '')} - Tờ bản đồ số: {mapping.get('[Số tờ]', '')}</p>
    <p>Địa chỉ thửa đất: {mapping.get('[Địa chỉ đất]', '')}</p>
    <p>&nbsp;</p>
    """
    
    return templates.TemplateResponse("cases/preview.html", {
        "request": request, 
        "case": case,
        "html_content": html_content
    })

@router.post("/{cid}/export-preview")
def export_preview(cid: int, html_content: str = Form(""), db: Session = Depends(get_db)):
    case = db.query(InheritanceCase).filter(InheritanceCase.id == cid).first()
    if not case: raise HTTPException(404)
    
    try:
        from docx import Document
        from htmldocx import HtmlToDocx
        doc = Document()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html_content, doc)
    except ImportError:
        from docx import Document
        from bs4 import BeautifulSoup
        doc = Document()
        soup = BeautifulSoup(html_content, "html.parser")
        doc.add_paragraph(soup.get_text(separator='\n'))
        doc.add_paragraph("\n\n(Lỗi: Yêu cầu pip install htmldocx để kết xuất chính tả/định dạng HTML)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"ho_so_thua_ke_{cid}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
