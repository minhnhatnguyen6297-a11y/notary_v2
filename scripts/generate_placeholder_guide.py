from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT_DIR / "word_templates" / "HUONG_DAN_PLACEHOLDER.docx"

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

PROPERTY_ROWS = [
    ("[Loại sổ]", "Loại Giấy chứng nhận", "Giấy chứng nhận quyền sử dụng đất"),
    ("[Serial]", "Số phát hành GCN", "BS 123456"),
    ("[Số vào sổ]", "Số vào sổ cấp GCN", "00123"),
    ("[Số thửa]", "Số thửa đất", "125"),
    ("[Số tờ]", "Số tờ bản đồ", "12"),
    ("[Địa chỉ đất]", "Địa chỉ thửa đất", "Xã Ninh Mỹ, huyện Hoa Lư, tỉnh Ninh Bình"),
    ("[Diện tích]", "Tổng diện tích (m², số)", "250"),
    ("[Diện tích chữ]", "Tổng diện tích bằng chữ", "Hai trăm năm mươi mét vuông"),
    ("[Hình thức sử dụng]", "Hình thức sử dụng đất", "Riêng"),
    ("[Loại đất]", "Loại đất tổng hợp trong văn bản", "Đất ở tại nông thôn"),
    ("[Nguồn gốc]", "Nguồn gốc sử dụng", "Nhà nước giao có thu tiền"),
    ("[Ngày cấp sổ]", "Ngày cấp GCN", "15/03/2010"),
    ("[Cơ quan cấp sổ]", "Cơ quan cấp GCN", "UBND huyện Hoa Lư"),
    ("[Loại đất N]", "Loại đất dòng N (N=1-10)", "Đất ở tại nông thôn"),
    ("[Diện tích N]", "Diện tích dòng N (N=1-10)", "120"),
    ("[Thời hạn N]", "Thời hạn sử dụng dòng N (N=1-10)", "Lâu dài"),
]

DATE_ROWS = [
    ("[Ngày]", "Ngày hiện tại (số)", "3"),
    ("[Tháng]", "Tháng hiện tại (2 chữ số)", "04"),
    ("[Ngày chữ]", "Ngày bằng chữ", "ba"),
    ("[Tháng chữ]", "Tháng bằng chữ", "tư"),
    ("[Niêm Yết]", "Xã/thị trấn nơi lập văn bản", "xã Ninh Mỹ"),
    ("[NIÊM YẾT]", "Như trên, IN HOA", "XÃ NINH MỸ"),
    ("[Tên file]", "Tên file gợi ý khi xuất văn bản", "ho_so_thua_ke_123"),
]

PERSON_ROWS = [
    ("[Xưng hô N]", "Danh xưng trong câu theo giới tính", "Ông / Bà"),
    ("[Tên N]", "Họ và tên đầy đủ", "Nguyễn Văn An"),
    ("[Năm sinh N]", "Ngày/năm sinh (01/01/YYYY -> chỉ in năm)", "1945"),
    ("[CCCD N]", "Số giấy tờ tùy thân", "036045001234"),
    ("[Loại CC N]", "Loại giấy tờ (tự tính từ ngày cấp)", "Căn cước công dân"),
    ("[Nơi cấp CC N]", "Cơ quan cấp (tự tính từ ngày cấp)", "Cục cảnh sát QLHC về TTXH"),
    ("[Ngày cấp N]", "Ngày cấp giấy tờ", "12/05/2020"),
    ("[Thường trú N]", "Cụm từ mở đầu địa chỉ (tự tính)", "Thường trú tại"),
    ("[Địa chỉ N]", "Địa chỉ thường trú/cư trú", "Xóm 3, xã Ninh Mỹ, huyện Hoa Lư"),
    ("[Năm chết N]", "Ngày mất (để trống nếu còn sống)", "20/01/2024"),
]

MANUAL_ROWS = [
    ("[Người ủy quyền]", "Tên người được ủy quyền ký", "Hệ thống để trống, người dùng tự điền"),
    ("[Người ủy quyền2]", "Người ủy quyền thứ 2", "Hệ thống để trống, người dùng tự điền"),
    ("[Số công chứng]", "Số văn bản chứng thực", "Hệ thống để trống, người dùng tự điền"),
    ("[SĐT]", "Số điện thoại", "Hệ thống để trống, người dùng tự điền"),
    ("[ONT]", "Diện tích đất ở nông thôn", "Hệ thống để trống, người dùng tự điền"),
    ("[CLN]", "Diện tích đất cây lâu năm", "Hệ thống để trống, người dùng tự điền"),
    ("[NTS]", "Diện tích đất nuôi trồng thủy sản", "Hệ thống để trống, người dùng tự điền"),
    ("[LUC]", "Diện tích đất lúa", "Hệ thống để trống, người dùng tự điền"),
    ("[Giá chuyển nhượng]", "Giá trị chuyển nhượng", "Hệ thống để trống, người dùng tự điền"),
]

NOTES = [
    "Placeholder phân biệt đúng chính tả và dấu ngoặc vuông. Ví dụ: [Niêm Yết] khác [NIÊM YẾT].",
    "Không thêm dấu cách thừa như [ Tên 1 ] vì hệ thống sẽ không thay được.",
    (
        "Word có thể tự tách placeholder khi định dạng. Nếu export ra mà token không được thay, "
        "hãy xóa token đó và gõ lại."
    ),
    "Không tự đặt tên placeholder mới. Chỉ dùng các placeholder có trong tài liệu này.",
]


def _set_default_style(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _style_run(run, font_name: str = "Times New Roman", size: float = 11.0, bold: bool = False,
               color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color is not None:
        run.font.color.rgb = color


def _set_cell_text(cell, text: str, *, font_name: str = "Times New Roman", size: float = 11.0,
                   bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    _style_run(run, font_name=font_name, size=size, bold=bold, color=color)


def _set_table_layout(table) -> None:
    table.style = "Table Grid"
    widths = (3.4, 7.0, 4.5)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] * 28.35)


def _add_placeholder_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=3)
    _set_table_layout(table)
    headers = ("Placeholder", "Ý nghĩa", "Ví dụ")
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, size=11.5)

    example_color = RGBColor(102, 102, 102)
    for row_index, (placeholder, meaning, example) in enumerate(rows, start=1):
        _set_cell_text(table.rows[row_index].cells[0], placeholder, font_name="Courier New", size=10.5)
        _set_cell_text(table.rows[row_index].cells[1], meaning, size=10.5)
        _set_cell_text(table.rows[row_index].cells[2], example, size=10.5, color=example_color)

    document.add_paragraph()


def main() -> int:
    document = Document()
    _set_default_style(document)

    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(title.add_run("HƯỚNG DẪN SỬ DỤNG PLACEHOLDER TRONG MẪU WORD"), size=16, bold=True)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _style_run(
        intro.add_run(
            "Placeholder là các nhãn hệ thống đặt sẵn trong file Word để khi xuất văn bản, dữ liệu hồ sơ sẽ được "
            "điền đúng vào vị trí tương ứng."
        ),
        size=12,
    )

    rule_heading = document.add_paragraph()
    _style_run(rule_heading.add_run("Quy tắc sử dụng"), size=13, bold=True)
    for bullet in (
        "Chỉ dùng các placeholder có trong tài liệu này, không tự đặt tên mới.",
        "Khi chỉnh sửa mẫu Word, người dùng chỉ nên thay đổi cách hành văn, xuống dòng, căn lề và bố cục.",
        "Placeholder phải được gõ đúng dấu ngoặc vuông, đúng chính tả và đúng khoảng trắng.",
    ):
        paragraph = document.add_paragraph(style="List Bullet")
        _style_run(paragraph.add_run(bullet), size=11.5)

    sample = document.add_paragraph()
    sample.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _style_run(
        sample.add_run(
            'Ví dụ: "[Xưng hô 1] [Tên 1] là con đẻ của [Xưng hô 2] [Tên 2]" là đúng; '
            '"[ Tên 1 ]" là sai vì có dấu cách thừa.'
        ),
        size=11.5,
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    _style_run(document.add_heading("Phần A - Thông tin tài sản đất", level=1).runs[0], size=14, bold=True)
    _add_placeholder_table(document, PROPERTY_ROWS)

    _style_run(document.add_heading("Phần B - Ngày tháng và thông tin hồ sơ", level=1).runs[0], size=14, bold=True)
    _add_placeholder_table(document, DATE_ROWS)

    person_heading = document.add_heading("Phần C - Thông tin người theo slot N", level=1)
    _style_run(person_heading.runs[0], size=14, bold=True)
    slot_note = document.add_paragraph()
    slot_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _style_run(
        slot_note.add_run(
            "Quy ước thứ tự: N=1 là người để lại di sản, N=2 là vợ/chồng, N=3 trở đi là người nhận hoặc từ chối "
            "theo sắp xếp của hệ thống."
        ),
        size=11.5,
    )
    _add_placeholder_table(document, PERSON_ROWS)

    _style_run(document.add_heading("Phần D - Placeholder điền tay", level=1).runs[0], size=14, bold=True)
    _add_placeholder_table(document, MANUAL_ROWS)

    _style_run(document.add_heading("Phần E - Lưu ý quan trọng", level=1).runs[0], size=14, bold=True)
    for note in NOTES:
        paragraph = document.add_paragraph(style="List Bullet")
        _style_run(paragraph.add_run(note), size=11.5)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Đã tạo file hướng dẫn: {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
