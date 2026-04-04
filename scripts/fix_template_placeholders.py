from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ROOT_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT_DIR / "word_templates" / "xa_PCDS_template.docx"
DEATH_PLACEHOLDER_PATTERN = re.compile(r"\[Năm chết\]")
HONORIFIC_PATTERN = re.compile(
    r"(?P<honorific>Ông/bà|ông/bà|Ông|Bà|ông|bà)(?P<separator>\s*:\s*|\s+)(?P<name>\[Tên (?P<slot>\d+)\])"
)

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass


def _iter_table_paragraphs(table: Table, prefix: str) -> Iterable[tuple[str, Paragraph]]:
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            location_prefix = f"{prefix} hàng {row_index} ô {cell_index}"
            yield from _iter_container_paragraphs(cell, location_prefix)


def _iter_header_footer_tables(section_part, prefix: str) -> Iterable[tuple[str, Paragraph]]:
    for paragraph_index, paragraph in enumerate(section_part.paragraphs, start=1):
        yield f"{prefix} đoạn {paragraph_index}", paragraph
    for table_index, table in enumerate(section_part.tables, start=1):
        yield from _iter_table_paragraphs(table, f"{prefix} bảng {table_index}")


def _iter_container_paragraphs(container: DocumentType | _Cell, prefix: str) -> Iterable[tuple[str, Paragraph]]:
    for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
        yield f"{prefix} đoạn {paragraph_index}", paragraph
    for table_index, table in enumerate(container.tables, start=1):
        yield from _iter_table_paragraphs(table, f"{prefix} bảng {table_index}")


def _iter_all_paragraphs(document: DocumentType) -> Iterable[tuple[str, Paragraph]]:
    yield from _iter_container_paragraphs(document, "thân tài liệu")
    for section_index, section in enumerate(document.sections, start=1):
        yield from _iter_header_footer_tables(section.header, f"section {section_index} header")
        yield from _iter_header_footer_tables(section.footer, f"section {section_index} footer")


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.clear()
        return
    paragraph.add_run(new_text)


def _transform_text(text: str) -> tuple[str, list[tuple[str, str, str]]]:
    changes: list[tuple[str, str, str]] = []

    def replace_death(match: re.Match[str]) -> str:
        old_value = match.group(0)
        new_value = "[Năm chết 1]"
        changes.append(("placeholder", old_value, new_value))
        return new_value

    updated_text = DEATH_PLACEHOLDER_PATTERN.sub(replace_death, text)

    def replace_honorific(match: re.Match[str]) -> str:
        slot = match.group("slot")
        separator = match.group("separator")
        name = match.group("name")
        old_value = match.group(0)
        new_value = f"[Xưng hô {slot}]{separator}{name}"
        changes.append(("xưng hô", old_value, new_value))
        return new_value

    updated_text = HONORIFIC_PATTERN.sub(replace_honorific, updated_text)
    return updated_text, changes


def main() -> int:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Không tìm thấy file mẫu: {TEMPLATE_PATH}")

    document = Document(str(TEMPLATE_PATH))
    total_changes = 0

    for location, paragraph in _iter_all_paragraphs(document):
        original_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        if not original_text.strip():
            continue

        updated_text, changes = _transform_text(original_text)
        if not changes or updated_text == original_text:
            continue

        _replace_paragraph_text(paragraph, updated_text)
        for change_type, old_value, new_value in changes:
            total_changes += 1
            print(f"[{change_type}] {location}: {old_value} -> {new_value}")

    document.save(str(TEMPLATE_PATH))
    print(f"Tổng số thay đổi: {total_changes}")
    print(f"Đã lưu file: {TEMPLATE_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
